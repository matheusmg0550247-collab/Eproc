import streamlit as st
import requests
import base64
from datetime import datetime, timedelta
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_secret(section: str, key: str) -> str:
    try:
        return st.secrets[section][key]
    except Exception:
        return ""

def get_brazil_time():
    return datetime.utcnow() - timedelta(hours=3)

def _send_webhook(url: str, payload: dict) -> bool:
    if not url:
        return False
    try:
        requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=5)
        return True
    except Exception as e:
        print(f"Erro ao enviar webhook: {e}")
        return False

def send_to_chat(webhook_key: str, text_msg: str) -> bool:
    """
    Envia mensagens para n8n quando configurado em [n8n] no secrets.toml.
    Mantém compatibilidade com chaves em [chat].

    Payload (padrão):
      { "tipo": "<webhook_key>", "text": "<mensagem>" }
    """
    try:
        n8n = st.secrets.get("n8n", {})
    except Exception:
        n8n = {}

    if webhook_key in ("bastao", "bastao_eq1", "bastao_eq2"):
        url = n8n.get("bastao_giro") or get_secret("chat", webhook_key) or get_secret("chat", "bastao")
    elif webhook_key in ("registro", "chamado", "checklist", "extras", "erro", "certidao"):
        url = n8n.get("registros") or get_secret("chat", webhook_key) or get_secret("chat", "registro")
    else:
        url = get_secret("chat", webhook_key)

    payload = {"tipo": webhook_key, "text": text_msg}
    return _send_webhook(url, payload)

def gerar_docx_certidao(tipo_certidao, num_processo, data_indisponibilidade_input, num_chamado, motivo_pedido):
    document = Document()
    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(12)

    head = document.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tj = head.add_run("TRIBUNAL DE JUSTIÇA DO ESTADO DE MINAS GERAIS\n")
    run_tj.bold = True

    num_parecer = int(datetime.now().strftime("%H%M"))
    ano_atual = datetime.now().year
    titulo = document.add_paragraph(f"Parecer Técnico GEJUD/DIRTEC/TJMG nº {num_parecer}/{ano_atual}.")
    titulo.runs[0].bold = True

    document.add_paragraph("Documento gerado automaticamente pelo Sistema de Controle de Bastão.")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

@st.cache_data
def get_img_as_base64(file_path: str):
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except Exception:
        return None
