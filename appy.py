# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
import requests
import time
from datetime import datetime, timedelta, date
from operator import itemgetter
from streamlit_autorefresh import st_autorefresh
import json
import re
import base64
import io
import altair as alt
from supabase import create_client
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ImportaÃ§Ã£o condicional
try:
    from streamlit_javascript import st_javascript
except ImportError:
    st_javascript = None

# ImportaÃ§Ãµes de utilitÃ¡rios
from utils import (get_brazil_time, get_secret, send_to_chat)

# ============================================
# 1. CONFIGURAÃ‡Ã•ES E CONSTANTES (EQUIPE ID 2 - CESUPE)
# ============================================
DB_APP_ID = 2        # ID da Fila desta equipe
LOGMEIN_DB_ID = 1    # ID do LogMeIn (COMPARTILHADO - SEMPRE 1)

CONSULTORES = sorted([
    "Barbara Mara", "Bruno Glaicon", "Claudia Luiza", "Douglas Paiva", "FÃ¡bio Alves", "Glayce Torres", 
    "Isabela Dias", "Isac Candido", "Ivana GuimarÃ£es", "Leonardo Damaceno", "Marcelo PenaGuerra", 
    "Michael Douglas", "MorÃ´ni", "Pablo Mol", "Ranyer Segal", "Sarah Leal", "Victoria Lisboa"
])

# Listas de OpÃ§Ãµes
REG_USUARIO_OPCOES = ["CartÃ³rio", "Gabinete", "Externo"]
REG_SISTEMA_OPCOES = ["Conveniados", "Outros", "Eproc", "Themis", "JPE", "SIAP"]
REG_CANAL_OPCOES = ["Presencial", "Telefone", "Email", "Whatsapp", "Outros"]
REG_DESFECHO_OPCOES = ["Resolvido - Cesupe", "Escalonado"]

OPCOES_ATIVIDADES_STATUS = ["HP", "E-mail", "WhatsApp PlantÃ£o", "HomologaÃ§Ã£o", "RedaÃ§Ã£o Documentos", "Outros"]

# URLs e Visuais (ATUALIZADO: FEVEREIRO LARANJA / CARNAVAL)
GIF_BASTAO_HOLDER = "https://media2.giphy.com/media/v1.Y2lkPTc5MGI3NjExa3Uwazd5cnNra2oxdDkydjZkcHdqcWN2cng0Y2N0cmNmN21vYXVzMiZlcD12MV9pbnRlcm5uYWxfZ2lmX2J5X2lkJmN0PWc/3rXs5J0hZkXwTZjuvM/giphy.gif"
GIF_LOGMEIN_TARGET = "https://media1.giphy.com/media/v1.Y2lkPTc5MGI3NjExZjFvczlzd3ExMWc2cWJrZ3EwNmplM285OGFqOHE1MXlzdnd4cndibiZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/mcsPU3SkKrYDdW3aAU/giphy.gif"
BASTAO_EMOJI = "ğŸ­" 
CARNAVAL_FILENAME = "Carnaval.gif"
APP_URL_CLOUD = 'https://controle-bastao-cesupe.streamlit.app'

# Cores Fevereiro Laranja
COR_LARANJA_PRIMARIA = "#FF8C00"
COR_LARANJA_SECUNDARIA = "#E65100"
COR_FUNDO_ALERTA = "#FFF3E0"

# Secrets
CHAT_WEBHOOK_BASTAO = get_secret("chat", "bastao")
WEBHOOK_STATE_DUMP = get_secret("webhook", "test_state")

# ============================================
# 2. OTIMIZAÃ‡ÃƒO E CONEXÃƒO
# ============================================

def get_supabase():
    try: 
        return create_client(st.secrets["supabase"]["url"], st.secrets["supabase"]["key"])
    except Exception as e:
        st.error(f"Erro ConexÃ£o DB: {e}") 
        return None

@st.cache_data(ttl=86400, show_spinner=False)
def carregar_dados_grafico():
    sb = get_supabase()
    if not sb: return None, None
    try:
        res = sb.table("atendimentos_resumo").select("data").eq("id", DB_APP_ID).execute()
        if res.data:
            json_data = res.data[0]['data']
            if 'totais_por_relatorio' in json_data:
                df = pd.DataFrame(json_data['totais_por_relatorio'])
                return df, json_data.get('gerado_em', '-')
    except Exception as e:
        st.error(f"Erro grÃ¡fico: {e}")
    return None, None

@st.cache_data
def get_img_as_base64_cached(file_path):
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except: return None

# ============================================
# 3. REPOSITÃ“RIO (LÃ“GICA ORIGINAL PRESERVADA)
# ============================================

def clean_data_for_db(obj):
    if isinstance(obj, dict):
        return {k: clean_data_for_db(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [clean_data_for_db(i) for i in obj]
    elif isinstance(obj, (datetime, date)):
        return obj.isoformat()
    elif isinstance(obj, timedelta):
        return obj.total_seconds()
    else:
        return obj

def load_state_from_db():
    sb = get_supabase()
    if not sb: return {}
    try:
        response = sb.table("app_state").select("data").eq("id", DB_APP_ID).execute()
        if response.data and len(response.data) > 0:
            return response.data[0].get("data", {})
        return {}
    except Exception as e:
        st.error(f"âš ï¸ Erro ao ler DB: {e}")
        return {}

def save_state_to_db(state_data):
    sb = get_supabase()
    if not sb: 
        st.error("Sem conexÃ£o para salvar.")
        return
    try:
        sanitized_data = clean_data_for_db(state_data)
        sb.table("app_state").upsert({"id": DB_APP_ID, "data": sanitized_data}).execute()
    except Exception as e:
        st.error(f"ğŸ”¥ ERRO DE ESCRITA NO BANCO: {e}")

def get_logmein_status():
    sb = get_supabase()
    if not sb: return None, False
    try:
        res = sb.table("controle_logmein").select("*").eq("id", LOGMEIN_DB_ID).execute()
        if res.data:
            return res.data[0].get('consultor_atual'), res.data[0].get('em_uso', False)
    except: pass
    return None, False

def set_logmein_status(consultor, em_uso):
    sb = get_supabase()
    if not sb: return
    try:
        dados = {
            "consultor_atual": consultor if em_uso else None,
            "em_uso": em_uso,
            "data_inicio": datetime.now().isoformat()
        }
        sb.table("controle_logmein").update(dados).eq("id", LOGMEIN_DB_ID).execute()
    except Exception as e: st.error(f"Erro LogMeIn DB: {e}")

# ============================================
# 4. FUNÃ‡Ã•ES DE UTILIDADE E IP
# ============================================
def get_browser_id():
    if st_javascript is None: return "no_js_lib"
    js_code = """(function() {
        let id = localStorage.getItem("device_id");
        if (!id) {
            id = "id_" + Math.random().toString(36).substr(2, 9);
            localStorage.setItem("device_id", id);
        }
        return id;
    })();"""
    try:
        return st_javascript(js_code, key="browser_id_tag")
    except: return "unknown_device"

def get_remote_ip():
    try:
        from streamlit.web.server.websocket_headers import ClientWebSocketRequest
        ctx = st.runtime.scriptrunner.get_script_run_ctx()
        if ctx and ctx.session_id:
            session_info = st.runtime.get_instance().get_client(ctx.session_id)
            if session_info:
                request = session_info.request
                if isinstance(request, ClientWebSocketRequest):
                    if 'X-Forwarded-For' in request.headers:
                        return request.headers['X-Forwarded-For'].split(',')[0]
                    return request.remote_ip
    except: return "Unknown"
    return "Unknown"

def get_ordered_visual_queue(queue, status_dict):
    if not queue: return []
    current_holder = next((c for c, s in status_dict.items() if 'BastÃ£o' in (s or '')), None)
    if not current_holder or current_holder not in queue: return list(queue)
    try:
        idx = queue.index(current_holder)
        return queue[idx:] + queue[:idx]
    except ValueError: return list(queue)

# ============================================
# 5. LÃ“GICA DE BANCO / CERTIDÃƒO / LOGS
# ============================================

def verificar_duplicidade_certidao(tipo, processo=None, data=None):
    sb = get_supabase()
    if not sb: return False
    try:
        query = sb.table("certidoes_registro").select("*")
        if tipo in ['FÃ­sico', 'EletrÃ´nico', 'FÃ­sica', 'EletrÃ´nica'] and processo:
            proc_limpo = str(processo).strip()
            if not proc_limpo: return False
            response = query.eq("processo", proc_limpo).execute()
            return len(response.data) > 0
        return False
    except Exception as e:
        print(f"Erro duplicidade: {e}")
        return False

def salvar_certidao_db(dados):
    sb = get_supabase()
    if not sb: return False
    try:
        if isinstance(dados.get('data'), (date, datetime)):
            dados['data'] = dados['data'].isoformat()
        if 'hora_periodo' in dados:
            if dados['hora_periodo']:
                dados['motivo'] = f"{dados.get('motivo', '')} - Hora/PerÃ­odo: {dados['hora_periodo']}"
            del dados['hora_periodo']
        if 'n_processo' in dados: 
            dados['processo'] = dados.pop('n_processo')
        if 'n_chamado' in dados: 
            dados['incidente'] = dados.pop('n_chamado')
        if 'data_evento' in dados:
            dados['data'] = dados.pop('data_evento')

        sb.table("certidoes_registro").insert(dados).execute()
        return True
    except Exception as e:
        st.error(f"Erro ao salvar no Supabase: {e}")
        return False

def gerar_docx_certidao_internal(tipo, numero, data, consultor, motivo, chamado="", hora="", nome_parte=""):
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0); section.right_margin = Cm(3.0)
        style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
        
        # CabeÃ§alho
        head_p = doc.add_paragraph(); head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runner = head_p.add_run("TRIBUNAL DE JUSTIÃ‡A DO ESTADO DE MINAS GERAIS\n")
        runner.bold = True
        head_p.add_run("Rua Ouro Preto, NÂ° 1564 - Bairro Santo Agostinho - CEP 30170-041 - Belo Horizonte - MG\nwww.tjmg.jus.br - Andar: 3Âº e 4Âº PV")
        doc.add_paragraph("\n")
        
        # NumeraÃ§Ã£o
        if tipo == 'Geral': p_num = doc.add_paragraph(f"Parecer GEJUD/DIRTEC/TJMG nÂº ____/2026. Assunto: Notifica erro no â€œJPe â€“ 2Âª InstÃ¢nciaâ€ ao peticionar.")
        else: p_num = doc.add_paragraph(f"Parecer TÃ©cnico GEJUD/DIRTEC/TJMG nÂº ____/2026. Assunto: Notifica erro no â€œJPe â€“ 2Âª InstÃ¢nciaâ€ ao peticionar.")
        p_num.runs[0].bold = True
        
        # Data por extenso
        data_extenso_str = ""
        try:
            dt_obj = datetime.strptime(data, "%d/%m/%Y")
            meses = {1:'janeiro', 2:'fevereiro', 3:'marÃ§o', 4:'abril', 5:'maio', 6:'junho', 
                     7:'julho', 8:'agosto', 9:'setembro', 10:'outubro', 11:'novembro', 12:'dezembro'}
            data_extenso_str = f"Belo Horizonte, {dt_obj.day} de {meses[dt_obj.month]} de {dt_obj.year}"
        except:
            data_extenso_str = f"Belo Horizonte, {data}" 
            
        doc.add_paragraph(data_extenso_str)
        doc.add_paragraph(f"Exmo(a). Senhor(a) Relator(a),")
        
        if tipo == 'Geral':
            corpo = doc.add_paragraph(); corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            txt = (f"Para fins de cumprimento dos artigos 13 e 14 da ResoluÃ§Ã£o nÂº 780/2014 do Tribunal de JustiÃ§a do Estado de Minas Gerais, "
                    f"informamos que em {data} houve indisponibilidade do portal JPe, superior a uma hora, {hora}, que impossibilitou o peticionamento eletrÃ´nico de recursos em processos que jÃ¡ tramitavam no sistema.")
            corpo.add_run(txt)
            doc.add_paragraph("\nColocamo-nos Ã  disposiÃ§Ã£o para outras que se fizerem necessÃ¡rias.")
            
        elif tipo in ['EletrÃ´nica', 'EletrÃ´nico']:
            corpo = doc.add_paragraph(); corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo.add_run(f"Informamos que de {data}, houve indisponibilidade especÃ­fica do sistema para o peticionamento do processo nÂº {numero}")
            if nome_parte: corpo.add_run(f", Parte/Advogado: {nome_parte}")
            corpo.add_run(".\n\n")
            corpo.add_run(f"O Chamado de nÃºmero {chamado if chamado else '_____'}, foi aberto e encaminhado Ã  DIRTEC (Diretoria Executiva de Tecnologia da InformaÃ§Ã£o e ComunicaÃ§Ã£o).\n\n")
            corpo.add_run("Esperamos ter prestado as informaÃ§Ãµes solicitadas e colocamo-nos Ã  disposiÃ§Ã£o para outras que se fizerem necessÃ¡rias.")

        elif tipo in ['FÃ­sica', 'FÃ­sico']:
            corpo1 = doc.add_paragraph(); corpo1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo1.add_run(f"Informamos que no dia {data}, houve indisponibilidade especÃ­fica do sistema para o peticionamento do processo nÂº {numero}")
            if nome_parte: corpo1.add_run(f", Parte/Advogado: {nome_parte}")
            corpo1.add_run(".")
            
            corpo2 = doc.add_paragraph(); corpo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo2.add_run(f"O Chamado de nÃºmero {chamado if chamado else '_____'}, foi aberto e encaminhado Ã  DIRTEC (Diretoria Executiva de Tecnologia da InformaÃ§Ã£o e ComunicaÃ§Ã£o).")
            
            corpo3 = doc.add_paragraph(); corpo3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo3.add_run("Diante da indisponibilidade especÃ­fica, nÃ£o havendo um prazo para soluÃ§Ã£o do problema, a Primeira Vice-PresidÃªncia recomenda o ingresso dos autos fÃ­sicos, nos termos do Â§ 2Âº, do artigo 14Âº, da ResoluÃ§Ã£o nÂº 780/2014, do Tribunal de JustiÃ§a do Estado de Minas Gerais.")
            doc.add_paragraph("Colocamo-nos Ã  disposiÃ§Ã£o para outras informaÃ§Ãµes que se fizerem necessÃ¡rias.")

        doc.add_paragraph("\nRespeitosamente,")
        sign = doc.add_paragraph("\n___________________________________\nWaner Andrade Silva\n0-009020-9\nCoordenaÃ§Ã£o de AnÃ¡lise e IntegraÃ§Ã£o de Sistemas Judiciais Informatizados - COJIN\nGerÃªncia de Sistemas Judiciais - GEJUD\nDiretoria Executiva de Tecnologia da InformaÃ§Ã£o e ComunicaÃ§Ã£o - DIRTEC")
        sign.runs[0].bold = True 
        
        buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
        return buffer
    except: return None

# ============================================
# 6. NOTIFICAÃ‡Ã•ES & WEBHOOKS
# ============================================
def send_chat_notification_internal(consultor, status):
    if CHAT_WEBHOOK_BASTAO and status == 'BastÃ£o':
        msg = f"ğŸ‰ **BASTÃƒO GIRADO!** ğŸ‰ \n\n- **Novo(a) ResponsÃ¡vel:** {consultor}\n- **Acesse o Painel:** {APP_URL_CLOUD}"
        try: send_to_chat("bastao", msg); return True
        except: return False
    return False

def send_horas_extras_to_chat(consultor, data, inicio, tempo, motivo):
    msg = f"â° **Registro de Horas Extras**\n\nğŸ‘¤ **Consultor:** {consultor}\nğŸ“… **Data:** {data.strftime('%d/%m/%Y')}\nğŸ• **InÃ­cio:** {inicio.strftime('%H:%M')}\nâ±ï¸ **Tempo Total:** {tempo}\nğŸ“ **Motivo:** {motivo}"
    try: send_to_chat("extras", msg); return True
    except: return False

def send_atendimento_to_chat(consultor, data, usuario, nome_setor, sistema, descricao, canal, desfecho, jira_opcional=""):
    jira_str = f"\nğŸ”¢ **Jira:** CESUPE-{jira_opcional}" if jira_opcional else ""
    msg = f"ğŸ“‹ **Novo Registro de Atendimento**\n\nğŸ‘¤ **Consultor:** {consultor}\nğŸ“… **Data:** {data.strftime('%d/%m/%Y')}\nğŸ‘¥ **UsuÃ¡rio:** {usuario}\nğŸ¢ **Nome/Setor:** {nome_setor}\nğŸ’» **Sistema:** {sistema}\nğŸ“ **DescriÃ§Ã£o:** {descricao}\nğŸ“ **Canal:** {canal}\nâœ… **Desfecho:** {desfecho}{jira_str}"
    try: send_to_chat("registro", msg); return True
    except: return False

def send_chamado_to_chat(consultor, texto):
    if not consultor or consultor == 'Selecione um nome' or not texto.strip(): return False
    data_envio = get_brazil_time().strftime('%d/%m/%Y %H:%M')
    msg = f"ğŸ†˜ **Rascunho de Chamado/Jira**\nğŸ“… **Data:** {data_envio}\n\nğŸ‘¤ **Autor:** {consultor}\n\nğŸ“ **Texto:**\n{texto}"
    try: send_to_chat('chamado', msg); return True
    except:
        try: send_to_chat('registro', msg); return True
        except: return False

def handle_erro_novidade_submission(consultor, titulo, objetivo, relato, resultado):
    data_envio = get_brazil_time().strftime("%d/%m/%Y %H:%M")
    msg = f"ğŸ› **Novo Relato de Erro/Novidade**\nğŸ“… **Data:** {data_envio}\n\nğŸ‘¤ **Autor:** {consultor}\nğŸ“Œ **TÃ­tulo:** {titulo}\n\nğŸ¯ **Objetivo:**\n{objetivo}\n\nğŸ§ª **Relato:**\n{relato}\n\nğŸ **Resultado:**\n{resultado}"
    try: send_to_chat("erro", msg); return True
    except: return False

# ============================================
# 7. GESTÃƒO DE ESTADO
# ============================================
def save_state():
    try:
        last_run = st.session_state.report_last_run_date
        visual_queue_calculated = get_ordered_visual_queue(st.session_state.bastao_queue, st.session_state.status_texto)
        
        state_to_save = {
            'status_texto': st.session_state.status_texto, 'bastao_queue': st.session_state.bastao_queue,
            'visual_queue': visual_queue_calculated, 'skip_flags': st.session_state.skip_flags, 
            'current_status_starts': st.session_state.current_status_starts,
            'bastao_counts': st.session_state.bastao_counts, 'priority_return_queue': st.session_state.priority_return_queue,
            'bastao_start_time': st.session_state.bastao_start_time, 
            'report_last_run_date': last_run, 
            'rotation_gif_start_time': st.session_state.get('rotation_gif_start_time'), 
            'auxilio_ativo': st.session_state.get('auxilio_ativo', False), 
            'daily_logs': st.session_state.daily_logs, 
            'simon_ranking': st.session_state.get('simon_ranking', []),
            'previous_states': st.session_state.get('previous_states', {})
        }
        save_state_to_db(state_to_save)
    except Exception as e: print(f"Erro save: {e}")

def format_time_duration(duration):
    if not isinstance(duration, timedelta): return '--:--:--'
    s = int(duration.total_seconds()); h, s = divmod(s, 3600); m, s = divmod(s, 60)
    return f'{h:02}:{m:02}:{s:02}'

def sync_state_from_db():
    try:
        db_data = load_state_from_db()
        if not db_data: return
        keys = ['status_texto', 'bastao_queue', 'skip_flags', 'bastao_counts', 'priority_return_queue', 'daily_logs', 'simon_ranking', 'previous_states']
        for k in keys:
            if k in db_data: st.session_state[k] = db_data[k]
        if 'bastao_start_time' in db_data and db_data['bastao_start_time']:
            try:
                if isinstance(db_data['bastao_start_time'], str): st.session_state['bastao_start_time'] = datetime.fromisoformat(db_data['bastao_start_time'])
                else: st.session_state['bastao_start_time'] = db_data['bastao_start_time']
            except: pass
        if 'current_status_starts' in db_data:
            starts = db_data['current_status_starts']
            for nome, val in starts.items():
                if isinstance(val, str):
                    try: st.session_state.current_status_starts[nome] = datetime.fromisoformat(val)
                    except: pass
                else: st.session_state.current_status_starts[nome] = val
    except Exception as e: print(f"Erro sync: {e}")

def log_status_change(consultor, old_status, new_status, duration):
    if not isinstance(duration, timedelta): duration = timedelta(0)
    now_br = get_brazil_time()
    old_lbl = old_status if old_status else 'Fila BastÃ£o'
    new_lbl = new_status if new_status else 'Fila BastÃ£o'
    if consultor in st.session_state.bastao_queue:
        if 'BastÃ£o' not in new_lbl and new_lbl != 'Fila BastÃ£o': new_lbl = f"Fila | {new_lbl}"
    
    device_id_audit = st.session_state.get('device_id_val', 'unknown')
    
    st.session_state.daily_logs.append({
        'timestamp': now_br, 'consultor': consultor, 
        'old_status': old_lbl, 'new_status': new_lbl, 
        'duration': duration, 'ip': device_id_audit
    })
    st.session_state.current_status_starts[consultor] = now_br

def update_status(novo_status: str, marcar_indisponivel: bool = False, manter_fila_atual: bool = False):
    selected = st.session_state.get('consultor_selectbox')
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    
    ensure_daily_reset()
    now_br = get_brazil_time()
    current = st.session_state.status_texto.get(selected, '')
    forced_successor = None
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in (s or '')), None)
    
    if novo_status == 'AlmoÃ§o':
        st.session_state.previous_states[selected] = {
            'status': current,
            'in_queue': selected in st.session_state.bastao_queue
        }
    
    if marcar_indisponivel:
        st.session_state[f'check_{selected}'] = False; st.session_state.skip_flags[selected] = True
        if selected in st.session_state.bastao_queue:
            st.session_state.bastao_queue.remove(selected)
            if selected not in st.session_state.priority_return_queue: st.session_state.priority_return_queue.append(selected)
    
    if novo_status == 'IndisponÃ­vel':
        if selected in st.session_state.bastao_queue:
            st.session_state.bastao_queue.remove(selected)
    
    elif manter_fila_atual:
        pass 
    
    clean_new = (novo_status or '').strip()
    if clean_new == 'Fila BastÃ£o': clean_new = ''
    
    final_status = clean_new
    if selected == current_holder and selected in st.session_state.bastao_queue:
         final_status = ('BastÃ£o | ' + clean_new).strip(' |') if clean_new else 'BastÃ£o'

    if not final_status and (selected not in st.session_state.bastao_queue): final_status = 'IndisponÃ­vel'
    
    try:
        started = st.session_state.current_status_starts.get(selected, now_br)
        log_status_change(selected, current, final_status, now_br - started)
    except: pass
    
    st.session_state.status_texto[selected] = final_status
    try: check_and_assume_baton(forced_successor=forced_successor)
    except: pass
    save_state()

def auto_manage_time():
    now = get_brazil_time(); last_run = st.session_state.report_last_run_date
    if now.hour >= 23 and now.date() == last_run.date(): reset_day_state(); save_state()
    elif now.date() > last_run.date(): reset_day_state(); save_state()

def find_next_holder_index(current_index, queue, skips):
    if not queue: return -1
    n = len(queue); start_index = (current_index + 1) % n
    for i in range(n):
        idx = (start_index + i) % n
        consultor = queue[idx]
        if not skips.get(consultor, False): return idx
    if n > 1:
        proximo_imediato_idx = (current_index + 1) % n
        nome_escolhido = queue[proximo_imediato_idx]; st.session_state.skip_flags[nome_escolhido] = False 
        return proximo_imediato_idx
    return -1

def check_and_assume_baton(forced_successor=None, immune_consultant=None):
    queue, skips = st.session_state.bastao_queue, st.session_state.skip_flags
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)
    is_valid = (current_holder and current_holder in queue)
    target = forced_successor if forced_successor else (current_holder if is_valid else None)
    
    if not target:
        curr_idx = queue.index(current_holder) if (current_holder and current_holder in queue) else -1
        idx = find_next_holder_index(curr_idx, queue, skips)
        target = queue[idx] if idx != -1 else None
        
    changed = False; now = get_brazil_time()
    for c in CONSULTORES:
        if c != immune_consultant: 
            if c != target and 'BastÃ£o' in st.session_state.status_texto.get(c, ''):
                log_status_change(c, 'BastÃ£o', 'IndisponÃ­vel', now - st.session_state.current_status_starts.get(c, now))
                st.session_state.status_texto[c] = 'IndisponÃ­vel'; changed = True
    
    if target:
        curr_s = st.session_state.status_texto.get(target, '')
        if 'BastÃ£o' not in curr_s:
            old_s = curr_s; new_s = f"BastÃ£o | {old_s}" if old_s and old_s != "IndisponÃ­vel" else "BastÃ£o"
            log_status_change(target, old_s, new_s, now - st.session_state.current_status_starts.get(target, now))
            st.session_state.status_texto[target] = new_s; st.session_state.bastao_start_time = now
            if current_holder != target: st.session_state.play_sound = True; send_chat_notification_internal(target, 'BastÃ£o')
            st.session_state.skip_flags[target] = False
            changed = True
    elif not target and current_holder:
        if current_holder != immune_consultant:
            log_status_change(current_holder, 'BastÃ£o', 'IndisponÃ­vel', now - st.session_state.current_status_starts.get(current_holder, now))
            st.session_state.status_texto[current_holder] = 'IndisponÃ­vel'; changed = True
            
    if changed: save_state()
    return changed

def init_session_state():
    dev_id = get_browser_id()
    if dev_id: st.session_state['device_id_val'] = dev_id

    if 'db_loaded' not in st.session_state:
        try:
            db_data = load_state_from_db()
            if db_data:
                for key, value in db_data.items(): st.session_state[key] = value
        except: pass
        st.session_state['db_loaded'] = True
    if 'report_last_run_date' in st.session_state and isinstance(st.session_state['report_last_run_date'], str):
        try: st.session_state['report_last_run_date'] = datetime.fromisoformat(st.session_state['report_last_run_date'])
        except: st.session_state['report_last_run_date'] = datetime.min
    now = get_brazil_time()
    defaults = {
        'bastao_start_time': None, 'report_last_run_date': datetime.min, 'rotation_gif_start_time': None,
        'play_sound': False, 'gif_warning': False, 'lunch_warning_info': None, 'last_reg_status': None,
        'chamado_guide_step': 0, 'auxilio_ativo': False, 'active_view': None, 'last_jira_number': "",
        'consultor_selectbox': "Selecione um nome", 'status_texto': {nome: 'IndisponÃ­vel' for nome in CONSULTORES},
        'bastao_queue': [], 'skip_flags': {}, 'current_status_starts': {nome: now for nome in CONSULTORES},
        'bastao_counts': {nome: 0 for nome in CONSULTORES}, 'priority_return_queue': [], 'daily_logs': [], 'simon_ranking': [],
        'word_buffer': None, 'aviso_duplicidade': False, 'previous_states': {}, 'view_logmein_ui': False
    }
    for key, default in defaults.items():
        if key not in st.session_state: st.session_state[key] = default
        
    for nome in CONSULTORES:
        st.session_state.bastao_counts.setdefault(nome, 0); st.session_state.skip_flags.setdefault(nome, False)
        current_status = st.session_state.status_texto.get(nome, 'IndisponÃ­vel')
        if current_status is None: current_status = 'IndisponÃ­vel'
        st.session_state.status_texto[nome] = current_status
        blocking = ['AlmoÃ§o', 'Ausente', 'SaÃ­da rÃ¡pida', 'SessÃ£o', 'ReuniÃ£o', 'Treinamento', 'Atendimento Presencial']
        is_hard_blocked = any(kw in current_status for kw in blocking)
        if is_hard_blocked: is_available = False
        elif nome in st.session_state.priority_return_queue: is_available = False
        elif nome in st.session_state.bastao_queue: is_available = True
        else: is_available = 'IndisponÃ­vel' not in current_status
        st.session_state[f'check_{nome}'] = is_available
        if nome not in st.session_state.current_status_starts: st.session_state.current_status_starts[nome] = now
    check_and_assume_baton()

def reset_day_state():
    now = get_brazil_time()
    st.session_state.bastao_queue = []; st.session_state.status_texto = {n: 'IndisponÃ­vel' for n in CONSULTORES}
    st.session_state.bastao_counts = {n: 0 for n in CONSULTORES}; st.session_state.skip_flags = {}
    st.session_state.daily_logs = []; st.session_state.current_status_starts = {n: now for n in CONSULTORES}
    st.session_state.report_last_run_date = now
    st.session_state.previous_states = {}
    for n in CONSULTORES: st.session_state[f'check_{n}'] = False

def ensure_daily_reset():
    now_br = get_brazil_time(); last_run = st.session_state.report_last_run_date
    if now_br.date() > last_run.date():
        reset_day_state(); st.toast("ğŸ§¹ Novo dia! Fila limpa."); save_state()

def toggle_queue(consultor):
    ensure_daily_reset(); st.session_state.gif_warning = False; now_br = get_brazil_time()
    if consultor in st.session_state.bastao_queue:
        current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)
        forced_successor = None
        if consultor == current_holder:
            idx = st.session_state.bastao_queue.index(consultor)
            nxt = find_next_holder_index(idx, st.session_state.bastao_queue, st.session_state.skip_flags)
            if nxt != -1: forced_successor = st.session_state.bastao_queue[nxt]
        st.session_state.bastao_queue.remove(consultor)
        st.session_state[f'check_{consultor}'] = False
        current_s = st.session_state.status_texto.get(consultor, '')
        if current_s == '' or current_s == 'BastÃ£o':
            log_status_change(consultor, current_s, 'IndisponÃ­vel', now_br - st.session_state.current_status_starts.get(consultor, now_br))
            st.session_state.status_texto[consultor] = 'IndisponÃ­vel'
        check_and_assume_baton(forced_successor)
    else:
        st.session_state.bastao_queue.append(consultor)
        st.session_state[f'check_{consultor}'] = True
        st.session_state.skip_flags[consultor] = False
        if consultor in st.session_state.priority_return_queue: st.session_state.priority_return_queue.remove(consultor)
        current_s = st.session_state.status_texto.get(consultor, 'IndisponÃ­vel')
        log_status_change(consultor, current_s, '', now_br - st.session_state.current_status_starts.get(consultor, now_br))
        st.session_state.status_texto[consultor] = ''
        check_and_assume_baton()
    save_state()

def leave_specific_status(consultor, status_type_to_remove):
    ensure_daily_reset(); st.session_state.gif_warning = False
    restored = False
    if consultor in st.session_state.previous_states:
        saved = st.session_state.previous_states.pop(consultor)
        if saved['in_queue'] and consultor not in st.session_state.bastao_queue: st.session_state.bastao_queue.append(consultor)
        elif not saved['in_queue'] and consultor in st.session_state.bastao_queue: st.session_state.bastao_queue.remove(consultor)
        old_status = st.session_state.status_texto.get(consultor, '')
        new_status = saved['status']
        st.session_state[f'check_{consultor}'] = (consultor in st.session_state.bastao_queue)
        st.session_state.skip_flags[consultor] = False
        restored = True
    else:
        if consultor not in st.session_state.bastao_queue: st.session_state.bastao_queue.append(consultor)
        st.session_state[f'check_{consultor}'] = True; st.session_state.skip_flags[consultor] = False
        new_status = '' 
    if not restored:
        old_status = st.session_state.status_texto.get(consultor, '')
        parts = [p.strip() for p in old_status.split('|')]
        new_parts = [p for p in parts if status_type_to_remove not in p and p]
        new_status = " | ".join(new_parts)
        if not new_status and consultor in st.session_state.bastao_queue: new_status = '' 
        elif not new_status: new_status = 'IndisponÃ­vel'
    now_br = get_brazil_time(); duration = now_br - st.session_state.current_status_starts.get(consultor, now_br)
    log_status_change(consultor, old_status, new_status, duration)
    st.session_state.status_texto[consultor] = new_status
    check_and_assume_baton(); save_state()

def enter_from_indisponivel(consultor):
    ensure_daily_reset(); st.session_state.gif_warning = False
    if consultor not in st.session_state.bastao_queue: st.session_state.bastao_queue.append(consultor)
    st.session_state[f'check_{consultor}'] = True; st.session_state.skip_flags[consultor] = False
    old_status = st.session_state.status_texto.get(consultor, 'IndisponÃ­vel')
    duration = get_brazil_time() - st.session_state.current_status_starts.get(consultor, get_brazil_time())
    log_status_change(consultor, old_status, '', duration)
    st.session_state.status_texto[consultor] = ''
    check_and_assume_baton(); save_state()

def rotate_bastao():
    ensure_daily_reset(); selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    queue = st.session_state.bastao_queue; skips = st.session_state.skip_flags
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)
    if selected != current_holder: st.error(f"âš ï¸ Apenas quem estÃ¡ com o bastÃ£o ({current_holder}) pode passÃ¡-lo!"); return
    current_index = queue.index(current_holder) if current_holder in queue else -1
    if current_index == -1: check_and_assume_baton(); return
    next_idx = find_next_holder_index(current_index, queue, skips)
    if next_idx != -1:
        next_holder = queue[next_idx]; now_br = get_brazil_time()
        old_h_status = st.session_state.status_texto[current_holder]
        new_h_status = old_h_status.replace('BastÃ£o | ', '').replace('BastÃ£o', '').strip()
        log_status_change(current_holder, old_h_status, new_h_status, now_br - (st.session_state.bastao_start_time or now_br))
        st.session_state.status_texto[current_holder] = new_h_status
        old_n_status = st.session_state.status_texto.get(next_holder, '')
        new_n_status = f"BastÃ£o | {old_n_status}" if old_n_status else "BastÃ£o"
        log_status_change(next_holder, old_n_status, new_n_status, timedelta(0))
        st.session_state.status_texto[next_holder] = new_n_status
        st.session_state.bastao_start_time = now_br
        st.session_state.bastao_counts[current_holder] = st.session_state.bastao_counts.get(current_holder, 0) + 1
        st.session_state.play_sound = True; send_chat_notification_internal(next_holder, 'BastÃ£o')
        save_state()
    else: st.warning('NinguÃ©m elegÃ­vel.'); check_and_assume_baton()

def toggle_view(v): 
    if st.session_state.active_view == v: st.session_state.active_view = None; return
    st.session_state.active_view = v
    if v == 'chamados': st.session_state.chamado_guide_step = 5 # DIRETO PARA TEXTO

def toggle_skip():
    selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um nome.'); return
    if selected not in st.session_state.bastao_queue: st.warning(f'{selected} nÃ£o estÃ¡ na fila.'); return
    st.session_state.skip_flags[selected] = not st.session_state.skip_flags.get(selected, False)
    save_state()

def toggle_presence_btn():
    selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um nome.'); return
    toggle_queue(selected)

# --- HANDLER CHAMADOS (COM TRAVA DE SEGURANÃ‡A) ---
def handle_chamado_submission():
    consultor = st.session_state.get('consultor_selectbox')
    texto = st.session_state.get('chamado_textarea', '')
    if not consultor or consultor == 'Selecione um nome':
        st.error("Selecione seu nome antes de enviar.")
        return False
    if not texto.strip():
        st.error("Erro ao enviar. Verifique o texto.")
        return False
    
    if send_chamado_to_chat(consultor, texto):
        st.session_state.chamado_textarea = ''; st.session_state.active_view = None
        st.toast("Enviado com sucesso!", icon="âœ…")
        return True
    return False

def open_logmein_ui(): st.session_state.view_logmein_ui = True
def close_logmein_ui(): st.session_state.view_logmein_ui = False

# ============================================
# 8. INTERFACE (FEVEREIRO LARANJA + CARNAVAL)
# ============================================
st.set_page_config(page_title="Controle Cesupe 2026", layout="wide", page_icon="ğŸ­")

# CSS para Laranja + Confetes
st.markdown(f"""
<style>
    div.stButton > button {{width: 100%; white-space: nowrap; height: 3rem;}} 
    [data-testid='stHorizontalBlock'] div.stButton > button {{white-space: nowrap; height: 3rem;}}
    
    /* AnimaÃ§Ã£o Confetes */
    @keyframes fall {{
        0% {{ top: -10px; transform: translateX(0) rotate(0deg); }}
        100% {{ top: 100vh; transform: translateX(50px) rotate(360deg); }}
    }}
    .confetti {{
        position: fixed; width: 10px; height: 10px;
        top: -10px; z-index: 9999; pointer-events: none;
        animation: fall 5s linear infinite;
    }}
</style>
<div class="confetti" style="left:10%; background-color:#FF8C00; animation-delay:0s;"></div>
<div class="confetti" style="left:25%; background-color:#FFD700; animation-delay:2s;"></div>
<div class="confetti" style="left:45%; background-color:#FF4500; animation-delay:1s;"></div>
<div class="confetti" style="left:65%; background-color:#FFA500; animation-delay:3s;"></div>
<div class="confetti" style="left:85%; background-color:#E65100; animation-delay:1.5s;"></div>
""", unsafe_allow_html=True)

# Bloco ConscientizaÃ§Ã£o Topo
st.markdown(f"""
<div style="background-color: {COR_FUNDO_ALERTA}; border-left: 6px solid {COR_LARANJA_PRIMARIA}; padding: 15px; border-radius: 8px; margin-bottom: 25px; box-shadow: 2px 2px 5px rgba(0,0,0,0.05);">
    <p style="margin: 0; color: #5D4037; font-size: 1.05rem; line-height: 1.5;">
    <b>Fevereiro Laranja</b> Ã© um convite Ã  consciÃªncia e Ã  aÃ§Ã£o: ele chama atenÃ§Ã£o para a leucemia e para a importÃ¢ncia do diagnÃ³stico precoce, que pode salvar vidas. ğŸ’›ğŸ§¡<br><br>
    Informar, apoiar quem estÃ¡ em tratamento e incentivar a doaÃ§Ã£o de sangue e de medula Ã³ssea sÃ£o atitudes que fazem diferenÃ§a. Compartilhe, converse e, se puder, cadastre-se como doador â€” um gesto simples pode ser a esperanÃ§a de alguÃ©m.
    </p>
</div>""", unsafe_allow_html=True)

init_session_state(); auto_manage_time()
st.components.v1.html("<script>window.scrollTo(0, 0);</script>", height=0)

# CabeÃ§alho Restaurado com Moldura Quadrada Maior
c_topo_esq, c_topo_dir = st.columns([2, 1], vertical_alignment="bottom")
with c_topo_esq:
    img_b64 = get_img_as_base64_cached(CARNAVAL_FILENAME)
    src = f"data:image/gif;base64,{img_b64}" if img_b64 else GIF_BASTAO_HOLDER
    st.markdown(f"""
    <div style="display: flex; align-items: center; gap: 25px;">
        <img src="{src}" style="width: 220px; height: 220px; border-radius: 15px; border: 5px solid {COR_LARANJA_PRIMARIA}; object-fit: cover; box-shadow: 0 4px 15px rgba(0,0,0,0.2);">
        <div>
            <span style="color: {COR_LARANJA_PRIMARIA}; font-weight: bold; letter-spacing: 2px; font-size: 1.1rem;">FEVEREIRO LARANJA</span>
            <h1 style="margin: 0; padding: 0; font-size: 2.3rem; color: {COR_LARANJA_SECUNDARIA}; text-shadow: 1px 1px 2px #AAA;">Controle BastÃ£o Cesupe 2026 {BASTAO_EMOJI}</h1>
        </div>
    </div>""", unsafe_allow_html=True)

with c_topo_dir:
    c_sub1, c_sub2 = st.columns([2, 1], vertical_alignment="bottom")
    with c_sub1: novo_responsavel = st.selectbox("Assumir BastÃ£o (RÃ¡pido)", options=["Selecione"] + CONSULTORES, label_visibility="collapsed", key="quick_enter")
    with c_sub2:
        if st.button("ğŸš€ Entrar", use_container_width=True):
            if novo_responsavel != "Selecione": toggle_queue(novo_responsavel); st.rerun()
    st.caption(f"ID SessÃ£o: ...{st.session_state.get('device_id_val', '???')[-4:]}")

st.markdown(f"<hr style='border: 1px solid {COR_LARANJA_PRIMARIA}; margin-top: 5px; margin-bottom: 20px;'>", unsafe_allow_html=True)

if st.session_state.active_view is None: st_autorefresh(interval=20000, key='auto_rerun'); sync_state_from_db() 
else: st.caption("â¸ï¸ AtualizaÃ§Ã£o automÃ¡tica pausada durante o registro.")

col_principal, col_disponibilidade = st.columns([1.5, 1])
queue, skips = st.session_state.bastao_queue, st.session_state.skip_flags
responsavel = next((c for c, s in st.session_state.status_texto.items() if 'BastÃ£o' in s), None)

with col_principal:
    st.header("ResponsÃ¡vel pelo BastÃ£o")
    if responsavel:
        st.markdown(f"""<div style="background: linear-gradient(135deg, {COR_FUNDO_ALERTA} 0%, #FFFFFF 100%); border: 3px solid {COR_LARANJA_PRIMARIA}; padding: 25px; border-radius: 15px; display: flex; align-items: center; box-shadow: 0 4px 15px rgba(255, 140, 0, 0.3); margin-bottom: 20px;"><div style="flex-shrink: 0; margin-right: 25px;"><img src="{GIF_BASTAO_HOLDER}" style="width: 90px; height: 90px; border-radius: 50%; object-fit: cover; border: 2px solid {COR_LARANJA_PRIMARIA};"></div><div><span style="font-size: 14px; color: #555; font-weight: bold; text-transform: uppercase; letter-spacing: 1.5px;">Atualmente com:</span><br><span style="font-size: 42px; font-weight: 800; color: #000080; line-height: 1.1;">{responsavel}</span></div></div>""", unsafe_allow_html=True)
        dur = get_brazil_time() - (st.session_state.bastao_start_time or get_brazil_time())
        st.caption(f"â±ï¸ Tempo com o bastÃ£o: **{format_time_duration(dur)}**")
    else: st.markdown('<h2>(NinguÃ©m com o bastÃ£o)</h2>', unsafe_allow_html=True)
    
    st.markdown("###"); st.header("**Consultor(a)**")
    c_nome, c_act1, c_act2, c_act3 = st.columns([2, 1, 1, 1], vertical_alignment="bottom")
    with c_nome: st.selectbox('Selecione:', ['Selecione um nome'] + CONSULTORES, key='consultor_selectbox', label_visibility='collapsed')
    with c_act1: st.button("ğŸ­ Fila", on_click=toggle_presence_btn, use_container_width=True)
    with c_act2: st.button('ğŸ¯ Passar', on_click=rotate_bastao, use_container_width=True)
    with c_act3: st.button('â­ï¸ Pular', on_click=toggle_skip, use_container_width=True)

    r2c1, r2c2, r2c3, r2c4, r2c5 = st.columns(5)
    r2c1.button('ğŸ“‹ Atividades', on_click=toggle_view, args=('menu_atividades',), use_container_width=True)
    r2c2.button('ğŸ—ï¸ Projeto', on_click=toggle_view, args=('menu_projetos',), use_container_width=True)
    r2c3.button('ğŸ“ Treinamento', on_click=toggle_view, args=('menu_treinamento',), use_container_width=True)
    r2c4.button('ğŸ“… ReuniÃ£o', on_click=toggle_view, args=('menu_reuniao',), use_container_width=True)
    r2c5.button('ğŸ½ï¸ AlmoÃ§o', on_click=update_status, args=('AlmoÃ§o', True), use_container_width=True)
    
    r3c1, r3c2, r3c3, r3c4 = st.columns(4)
    r3c1.button('ğŸ™ï¸ SessÃ£o', on_click=toggle_view, args=('menu_sessao',), use_container_width=True)
    r3c2.button('ğŸš¶ SaÃ­da', on_click=update_status, args=('SaÃ­da rÃ¡pida', True), use_container_width=True)
    r3c3.button('ğŸƒ Sair', on_click=update_status, args=('IndisponÃ­vel', True), use_container_width=True)
    if r3c4.button("ğŸ¤ Atend. Presencial", use_container_width=True): toggle_view('menu_presencial')

    # MENUS DE AÃ‡ÃƒO (LÃ“GICA ORIGINAL)
    if st.session_state.active_view == 'menu_atividades':
        with st.container(border=True):
            at_t = st.multiselect("Tipo:", OPCOES_ATIVIDADES_STATUS); at_e = st.text_input("Detalhe:")
            if st.button("Confirmar Atividade", type="primary"): 
                update_status(f"Atividade: {', '.join(at_t)} - {at_e}", manter_fila_atual=True)
                st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == 'menu_presencial':
        with st.container(border=True):
            st.subheader('ğŸ¤ Registro Presencial'); l_p = st.text_input('Local:'); o_p = st.text_input('Objetivo:')
            if st.button('âœ… Confirmar', type='primary'):
                update_status(f"Atendimento Presencial: {l_p} - {o_p}", True)
                st.session_state.active_view = None; st.rerun()

    # LogMeIn UI Completa
    if st.button('ğŸ”‘ LogMeIn', use_container_width=True): open_logmein_ui()
    if st.session_state.get('view_logmein_ui'):
        with st.container(border=True):
            st.markdown("### ğŸ’» Acesso LogMeIn")
            l_user, l_in_use = get_logmein_status()
            if l_in_use: st.error(f"ğŸ”´ EM USO POR: {l_user}")
            else: st.success("âœ… DISPONÃVEL")
            if st.button("Fechar LogMeIn"): close_logmein_ui(); st.rerun()

    st.markdown("---")
    # FERRAMENTAS (TODOS OS 7 BOTÃ•ES RESTAURADOS)
    c_tool1, c_tool2, c_tool3, c_tool4, c_tool5, c_tool6, c_tool7 = st.columns(7)
    c_tool1.button("ğŸ“‘ Checklist", use_container_width=True, on_click=toggle_view, args=("checklist",))
    c_tool2.button("ğŸ†˜ Chamados", use_container_width=True, on_click=toggle_view, args=("chamados",))
    c_tool3.button("ğŸ“ Atend.", use_container_width=True, on_click=toggle_view, args=("atendimentos",))
    c_tool4.button("â° H. Extras", use_container_width=True, on_click=toggle_view, args=("hextras",))
    c_tool5.button("ğŸ› Erro/Nov.", use_container_width=True, on_click=toggle_view, args=("erro_novidade",))
    c_tool6.button("ğŸ–¨ï¸ CertidÃ£o", use_container_width=True, on_click=toggle_view, args=("certidao",))
    c_tool7.button("ğŸ’¡ SugestÃ£o", use_container_width=True, on_click=toggle_view, args=("sugestao",))

    if st.session_state.active_view == "chamados":
        with st.container(border=True):
            st.header('ğŸ†˜ Registro de Chamado (Jira)')
            st.text_area('Descreva o incidente:', height=240, key='chamado_textarea')
            if st.button('Enviar', type='primary', use_container_width=True):
                if handle_chamado_submission(): st.rerun()

    if st.session_state.active_view == "atendimentos":
        with st.container(border=True):
            at_d = st.date_input('Data:', value=get_brazil_time().date())
            at_u = st.selectbox('UsuÃ¡rio:', REG_USUARIO_OPCOES); at_s = st.text_input('Setor:')
            at_desc = st.text_area('DescriÃ§Ã£o:')
            if st.button('Enviar Atendimento', type='primary'):
                if send_atendimento_to_chat(st.session_state.consultor_selectbox, at_d, at_u, at_s, "Eproc", at_desc, "Email", "Resolvido"):
                    st.success('Enviado!'); st.session_state.active_view = None; st.rerun()

    if st.session_state.active_view == "certidao":
        with st.container(border=True):
            st.header("ğŸ–¨ï¸ CertidÃ£o (2026)"); c_dt = st.date_input("Data Evento:"); c_t = st.selectbox("Tipo:", ["EletrÃ´nica", "FÃ­sica", "Geral"])
            c_p = st.text_input("Processo:"); c_c = st.text_input("Chamado:"); c_m = st.text_area("Motivo:")
            c1, c2 = st.columns(2)
            with c1:
                if st.button("ğŸ“„ Gerar"): 
                    st.session_state.word_buffer = gerar_docx_certidao_internal(c_t, c_p, c_dt.strftime("%d/%m/%Y"), st.session_state.consultor_selectbox, c_m, c_c)
                if st.session_state.word_buffer: st.download_button("â¬‡ï¸ Baixar", st.session_state.word_buffer, file_name="certidao.docx")
            with c2:
                if st.button("ğŸ’¾ Salvar DB", type="primary"):
                    payload = {"tipo": c_t, "data": c_dt.isoformat(), "consultor": st.session_state.consultor_selectbox, "incidente": c_c, "processo": c_p, "motivo": c_m}
                    if salvar_certidao_db(payload): st.success("Salvo!"); time.sleep(1); st.session_state.active_view = None; st.rerun()

with col_disponibilidade:
    st.header('Status dos Consultores')
    # Listas de Status Completas
    ui_lists = {'fila': [], 'almoco': [], 'saida': [], 'ausente': [], 'atividade': [], 'sessao': [], 'projeto': [], 'reuniao': [], 'treinamento': [], 'indisponivel': [], 'presencial': []}
    for nome in CONSULTORES:
        if nome in queue: ui_lists['fila'].append(nome)
        s = st.session_state.status_texto.get(nome, 'IndisponÃ­vel')
        if s == 'AlmoÃ§o': ui_lists['almoco'].append(nome)
        elif s == 'SaÃ­da rÃ¡pida': ui_lists['saida'].append(nome)
        elif 'SessÃ£o:' in str(s): ui_lists['sessao'].append((nome, s.replace('SessÃ£o:', '')))
        elif 'Projeto:' in str(s): ui_lists['projeto'].append((nome, s.replace('Projeto:', '')))
        elif 'Atividade:' in str(s): ui_lists['atividade'].append((nome, s.replace('Atividade:', '')))
        elif 'ReuniÃ£o:' in str(s): ui_lists['reuniao'].append((nome, s.replace('ReuniÃ£o:', '')))
        elif 'Treinamento:' in str(s): ui_lists['treinamento'].append((nome, s.replace('Treinamento:', '')))
        elif 'Atendimento Presencial:' in str(s): ui_lists['presencial'].append((nome, s.replace('Atendimento Presencial:', '')))
        elif s == 'IndisponÃ­vel' and nome not in queue: ui_lists['indisponivel'].append(nome)

    st.subheader(f'âœ… Na Fila ({len(ui_lists["fila"])})')
    ordem = get_ordered_visual_queue(queue, st.session_state.status_texto)
    for i, n in enumerate(ordem):
        if n in ui_lists['fila']:
            clr = COR_LARANJA_PRIMARIA if n == responsavel else COR_FUNDO_ALERTA
            t_clr = "#FFF" if n == responsavel else "#333"
            st.markdown(f"<div style='background:{clr}; color:{t_clr}; padding:10px; border-radius:8px; margin-bottom:5px;'><b>{i+1}Âº {n}</b></div>", unsafe_allow_html=True)

    st.markdown("---")
    def _render_section(titulo, icon, itens):
        st.subheader(f'{icon} {titulo} ({len(itens)})')
        if not itens: st.markdown('_Nenhum._')
        else:
            for it in itens:
                n = it[0] if isinstance(it, tuple) else it
                d = it[1] if isinstance(it, tuple) else titulo
                st.markdown(f"<div style='background:{COR_FUNDO_ALERTA}; padding:8px; border-radius:8px; margin-bottom:4px;'><b>{n}</b>: <small>{d}</small></div>", unsafe_allow_html=True)

    _render_section('Atend. Presencial', 'ğŸ¤', ui_lists['presencial'])
    _render_section('Em Demanda', 'ğŸ“‹', ui_lists['atividade'])
    _render_section('Projetos', 'ğŸ—ï¸', ui_lists['projeto'])
    _render_section('Treinamento', 'ğŸ“', ui_lists['treinamento'])
    _render_section('ReuniÃ£o', 'ğŸ“…', ui_lists['reuniao'])
    _render_section('AlmoÃ§o', 'ğŸ½ï¸', ui_lists['almoco'])
    _render_section('SessÃ£o', 'ğŸ™ï¸', ui_lists['sessao'])
    _render_section('IndisponÃ­vel', 'âŒ', ui_lists['indisponivel'])
