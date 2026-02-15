# -*- coding: utf-8 -*-
import streamlit as st

# Auto-refresh (opcional)
try:
    from streamlit_autorefresh import st_autorefresh
except Exception:
    st_autorefresh = None
import requests
import time
import gc
from datetime import datetime, timedelta, date
from operator import itemgetter
import json
import hashlib
import re
import uuid
import unicodedata
import base64
import io
from supabase import create_client
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Importação condicional
try:
    from streamlit_javascript import st_javascript
except ImportError:
    st_javascript = None

# Importações de utilitários
from utils import (get_brazil_time, get_secret, send_to_chat)

# ============================================
# 1. CONFIGURAÇÕES E CONSTANTES (EQUIPE ID 2 - CESUPE)
# ============================================
DB_APP_ID = 2        # ID da Fila desta equipe
LOGMEIN_DB_ID = 1    # ID do LogMeIn (COMPARTILHADO - SEMPRE 1)
GLOBAL_STATE_VERSION_ID = 9999  # Carimbo global (qualquer equipe) para disparar refresh

CONSULTORES = sorted([
    "Barbara Mara", "Bruno Glaicon", "Claudia Luiza", "Douglas Paiva", "Fábio Alves", "Glayce Torres", 
    "Isabela Dias", "Isac Candido", "Ivana Guimarães", "Leonardo Damaceno", "Marcelo PenaGuerra", 
    "Michael Douglas", "Morôni", "Pablo Mol", "Ranyer Segal", "Sarah Leal", "Victoria Lisboa"
])

# Listas de Opções
REG_USUARIO_OPCOES = ["Cartório", "Gabinete", "Externo"]
REG_SISTEMA_OPCOES = ["Conveniados", "Outros", "Eproc", "Themis", "JPE", "SIAP"]
REG_CANAL_OPCOES = ["Presencial", "Telefone", "Email", "Whatsapp", "Outros"]
REG_DESFECHO_OPCOES = ["Resolvido - Cesupe", "Escalonado"]

OPCOES_ATIVIDADES_STATUS = ["HP", "E-mail", "WhatsApp Plantão", "Homologação", "Redação Documentos", "Outros"]

# URLs e Visuais
GIF_BASTAO_HOLDER = "https://media2.giphy.com/media/v1.Y2lkPTc5MGI3NjExa3Uwazd5cnNra2oxdDkydjZkcHdqcWN2cng0Y2N0cmNmN21vYXVzMiZlcD12MV9pbnRlcm5uYWxfZ2lmX2J5X2lkJmN0PWc/3rXs5J0hZkXwTZjuvM/giphy.gif"
GIF_LOGMEIN_TARGET = "https://media1.giphy.com/media/v1.Y2lkPTc5MGI3NjExZjFvczlzd3ExMWc2cWJrZ3EwNmplM285OGFqOHE1MXlzdnd4cndibiZlcD12MV9pbnRlcm5hbF9naWZfYnlfaWQmY3Q9Zw/mcsPU3SkKrYDdW3aAU/giphy.gif"
BASTAO_EMOJI = "🎭" 
PUG2026_FILENAME = "Carnaval.gif" 

# ============================================
# RAMAIS CESUPE (BADGE AO LADO DO NOME)
# ============================================
RAMAIS_CESUPE = {
    # Chave normalizada (sem acento, minúscula). Priorize combinações específicas antes do primeiro nome.
    'douglas paiva': '2663',
    'alex': '2510',
    'barbara': '2517',
    'bruno': '2644',
    'claudio': '2667',
    'claudia': '2667',  # variação comum
    'dirceu': '2666',
    'douglas': '2659',
    'fabio': '2665',
    'farley': '2651',
    'gilberto': '2654',
    'gleis': '2536',
    'gleissiane': '2536',
    'gleyce': '2647',
    'glayce': '2647',  # variação comum
    'hugo': '2650',
    'jerry': '2654',
    'jonatas': '2656',
    'leandro': '2652',
    'leonardo': '2655',
    'ivana': '2653',
    'marcelo': '2655',
    'matheus': '2664',
    'michael': '2638',
    'pablo': '2643',
    'ranier': '2669',
    'ranyer': '2669',  # variação comum
    'vanessa': '2607',
    'victoria': '2660',
    'victória': '2660',  # caso venha com acento
}


# ============================================
# AGENDA EPROC (VISÃO SEMANAL CONSOLIDADA)
# Fonte: "Distribuição das Atividades da Equipe EPROC" (anexo do usuário)
# ============================================
EPROC_VISAO_SEMANAL = [
    {
        "dia": "Segunda-feira",
        "manha": "Fábio, Leonardo, Glayce, Pablo",
        "tarde": "Fábio, Leonardo, Glayce, Isabela (Bruno, Cláudia, Douglas – Cond.)",
        "sessao": "Consultores sem projeto fixo e Configuração EPROC",
        "obs": ""
    },
    {
        "dia": "Terça-feira",
        "manha": "Fábio, Leonardo, Glayce, Pablo",
        "tarde": "Fábio, Leonardo, Glayce, Isabela (IA / SOMA / Manuais – Cond.)",
        "sessao": "Prioridade: Isac e Ranyer (SOMA)",
        "obs": ""
    },
    {
        "dia": "Quarta-feira",
        "manha": "Fábio, Leonardo, Glayce, Pablo",
        "tarde": "Fábio, Leonardo, Glayce, Isabela (IA / Manuais / Cartórios – Cond.)",
        "sessao": "Distribuição geral, sem prioridade específica",
        "obs": ""
    },
    {
        "dia": "Quinta-feira",
        "manha": "Fábio, Leonardo, Glayce, Pablo",
        "tarde": "Fábio, Leonardo, Glayce, Isabela (IA / SOMA / Manuais – Cond.)",
        "sessao": "Prioridade: Isac e Ranyer (SOMA)",
        "obs": "Obs.: Bárbara não pode."
    },
    {
        "dia": "Sexta-feira",
        "manha": "Fábio, Leonardo, Glayce (Demais – Cond.)",
        "tarde": "Fábio, Leonardo, Glayce (Demais – Cond.)",
        "sessao": "Preferencialmente consultores sem projeto crítico",
        "obs": ""
    },
]

DIA_SEMANA_PT = {
    0: "Segunda-feira",
    1: "Terça-feira",
    2: "Quarta-feira",
    3: "Quinta-feira",
    4: "Sexta-feira",
    5: "Sábado",
    6: "Domingo",
}

def render_agenda_eproc_sidebar():
    """Renderiza a agenda EPROC (visão semanal) no painel lateral."""
    now_br = get_brazil_time()
    dia_pt = DIA_SEMANA_PT.get(now_br.weekday(), "")
    st.markdown(f"### 📅 Agenda EPROC")
    st.caption(f"Hoje: **{dia_pt}** — {now_br.strftime('%d/%m/%Y')}")

    with st.expander("Ver visão semanal consolidada", expanded=False):
        for row in EPROC_VISAO_SEMANAL:
            is_today = (row.get("dia") == dia_pt)
            bg = "#E7F1FF" if is_today else "#FFFFFF"
            border = "1px solid #cbd5e1" if is_today else "1px solid #eee"
            
            # --- CORREÇÃO 1: Construção segura do HTML para evitar erro de sintaxe ---
            obs_text = row.get('obs')
            obs_html = f"<div style='font-size:13px; margin-top:6px; color:#6b7280;'><i>{obs_text}</i></div>" if obs_text else ""
            
            st.markdown(
                f"""
<div style='border:{border}; background:{bg}; padding:10px 12px; border-radius:12px; margin:8px 0;'>
  <div style='font-weight:800; font-size:15px; margin-bottom:6px;'>🗓️ {row.get('dia','')}</div>
  <div style='font-size:14px; margin:2px 0;'><b>🕘 Manhã:</b> {row.get('manha','')}</div>
  <div style='font-size:14px; margin:2px 0;'><b>🕜 Tarde:</b> {row.get('tarde','')}</div>
  <div style='font-size:14px; margin:2px 0;'><b>🎙️ Sessões:</b> {row.get('sessao','')}</div>
  {obs_html}
</div>
""", unsafe_allow_html=True
            )

def _normalize_nome(txt: str) -> str:
    if not isinstance(txt, str):
        return ''
    # remove acentos + normaliza espaços
    txt = ''.join(ch for ch in unicodedata.normalize('NFKD', txt) if not unicodedata.combining(ch))
    txt = re.sub(r'\s+', ' ', txt).strip().lower()
    return txt

def get_ramal_nome(nome: str):
    n = _normalize_nome(nome)
    if not n:
        return None
    # Regra específica: Douglas Paiva
    if 'douglas' in n and 'paiva' in n:
        return RAMAIS_CESUPE.get('douglas paiva')
    # Primeira tentativa: match completo
    if n in RAMAIS_CESUPE:
        return RAMAIS_CESUPE[n]
    # Segunda tentativa: primeiro nome
    first = n.split(' ')[0]
    return RAMAIS_CESUPE.get(first)

def _badge_ramal_html(ramal):
    if not ramal:
        return ''
    return f"<span style='margin-left:8px; padding:2px 8px; border-radius:999px; border:1px solid #ddd; font-size:12px; background:#f7f7f7;'>☎ {ramal}</span>"

def _icons_telefone_cafe(indic: dict):
    if not isinstance(indic, dict):
        return ''
    parts = []
    if indic.get('telefone'): parts.append('📞')
    if indic.get('cafe'): parts.append('☕')
    return (' ' + ' '.join(parts)) if parts else ''
APP_URL_CLOUD = 'https://controle-bastao-cesupe.streamlit.app'

# Secrets
# N8N / WhatsApp (Z-API via n8n)
# Preferencialmente configure em .streamlit/secrets.toml:
# [n8n]
# bastao_giro = "https://.../webhook/...."
# registros   = "https://.../webhook/...."
#
# Compatibilidade: caso não exista [n8n], usamos as chaves antigas em [chat]
# - chat.bastao_eq1 / chat.bastao_eq2 (giro do bastão)
# - chat.registro (demais registros/ferramentas)
N8N_WEBHOOK_BASTAO_GIRO = get_secret("n8n", "bastao_giro") or get_secret("chat", "bastao_eq1") or get_secret("chat", "bastao_eq2")
N8N_WEBHOOK_REGISTROS   = get_secret("n8n", "registros")   or get_secret("chat", "registro")

WEBHOOK_STATE_DUMP = get_secret("webhook", "test_state")

def post_n8n(url: str, payload: dict) -> bool:
    """Envia evento para n8n (silencioso). Retorna True/False."""
    if not url:
        return False
    try:
        # Timeout maior para evitar falhas intermitentes no Streamlit Cloud / DO
        resp = requests.post(url, json=payload, timeout=12)
        # Se der erro HTTP, considera falha (mas sem estourar exception pro usuário)
        if getattr(resp, "status_code", 200) >= 400:
            return False
        return True
    except requests.exceptions.Timeout:
        return False
    except Exception:
        return False

# ============================================
# 2. OTIMIZAÇÃO E CONEXÃO
# ============================================

@st.cache_resource(ttl=3600)
def get_supabase():
    try: 
        return create_client(st.secrets["supabase"]["url"], st.secrets["supabase"]["key"])
    except Exception as e:
        st.cache_resource.clear()
        return None

def render_operational_summary():
    """Resumo Operacional: link externo (sem gráfico no app)."""
    st.subheader("📊 Resumo Operacional")
    st.markdown('🔗 **Dados Semanais Cesupe** — [Abrir no Looker Studio](https://lookerstudio.google.com/reporting/8adc4cf8-94f6-4333-b808-4a300151af09)', unsafe_allow_html=False)


def get_img_as_base64_cached(file_path):
    try:
        with open(file_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    except: return None

# ============================================
# 3. REPOSITÓRIO
# ============================================

def clean_data_for_db(obj):
    if isinstance(obj, dict):
        return {k: clean_data_for_db(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [clean_data_for_db(i) for i in obj]
    elif isinstance(obj, (datetime, date)):
        return obj.isoformat()
    elif isinstance(obj, timedelta):
        return obj.total_seconds()
    else:
        return obj

@st.cache_data(ttl=5, show_spinner=False)
def load_state_from_db(app_id):
    sb = get_supabase()
    if not sb: return {}
    try:
        response = sb.table("app_state").select("data").eq("id", app_id).execute()
        if response.data and len(response.data) > 0:
            return response.data[0].get("data", {})
        return {}
    except Exception as e:
        return {}


@st.cache_data(ttl=2, show_spinner=False)
def load_global_state_version():
    """Lê o carimbo de versão global (qualquer equipe)."""
    sb = get_supabase()
    if not sb:
        return 0
    try:
        res = sb.table("app_state").select("data").eq("id", GLOBAL_STATE_VERSION_ID).execute()
        if res.data and len(res.data) > 0:
            data = res.data[0].get("data", {}) or {}
            try:
                return int(data.get("state_version") or 0)
            except Exception:
                return 0
        return 0
    except Exception:
        return 0

def bump_global_state_version(sb=None):
    """Atualiza o carimbo global para que todos os clientes percebam mudanças."""
    try:
        _sb = sb or get_supabase()
        if not _sb:
            return
        ver = int(time.time() * 1000)
        _sb.table("app_state").upsert({"id": GLOBAL_STATE_VERSION_ID, "data": {"state_version": ver}}).execute()
        try:
            load_global_state_version.clear()
        except Exception:
            pass
        st.session_state["_global_state_version_seen"] = ver
    except Exception:
        pass

def save_state_to_db(app_id, state_data):
    sb = get_supabase()
    if not sb: 
        st.error("Sem conexão para salvar.")
        return
    try:
        sanitized_data = clean_data_for_db(state_data)
        sb.table("app_state").upsert({"id": app_id, "data": sanitized_data}).execute()
        # Dispara refresh global para todos os clientes (qualquer equipe)
        bump_global_state_version(sb)

        # --- CORREÇÃO 2: Marca o tempo do salvamento para o Watcher não conflitar
        st.session_state['_last_save_time'] = time.time() 
        
        # Feedback visual pós-salvar (mostra no próximo rerun)
        try:
            st.session_state['_toast_msg'] = '✅ Registro salvo.'
        except Exception:
            pass
    except Exception as e:
        st.error(f"🔥 ERRO DE ESCRITA NO BANCO: {e}")

# --- LOGMEIN DB ---
def get_logmein_status():
    sb = get_supabase()
    if not sb: return None, False
    try:
        res = sb.table("controle_logmein").select("*").eq("id", LOGMEIN_DB_ID).execute()
        if res.data:
            return res.data[0].get('consultor_atual'), res.data[0].get('em_uso', False)
    except: pass
    return None, False

def set_logmein_status(consultor, em_uso):
    sb = get_supabase()
    if not sb: return
    try:
        dados = {
            "consultor_atual": consultor if em_uso else None,
            "em_uso": em_uso,
            "data_inicio": datetime.now().isoformat()
        }
        sb.table("controle_logmein").update(dados).eq("id", LOGMEIN_DB_ID).execute()
    except Exception as e: st.error(f"Erro LogMeIn DB: {e}")

# ============================================
# 4. FUNÇÕES DE UTILIDADE E IP
# ============================================
def get_browser_id():
    if st_javascript is None: return "no_js_lib"
    js_code = """(function() {
        let id = localStorage.getItem("device_id");
        if (!id) {
            id = "id_" + Math.random().toString(36).substr(2, 9);
            localStorage.setItem("device_id", id);
        }
        return id;
    })();"""
    try:
        return st_javascript(js_code, key="browser_id_tag")
    except: return "unknown_device"

def get_remote_ip():
    try:
        from streamlit.web.server.websocket_headers import ClientWebSocketRequest
        ctx = st.runtime.scriptrunner.get_script_run_ctx()
        if ctx and ctx.session_id:
            session_info = st.runtime.get_instance().get_client(ctx.session_id)
            if session_info:
                request = session_info.request
                if isinstance(request, ClientWebSocketRequest):
                    if 'X-Forwarded-For' in request.headers:
                        return request.headers['X-Forwarded-For'].split(',')[0]
                    return request.remote_ip
    except: return "Unknown"
    return "Unknown"

# --- LIMPEZA DE MEMÓRIA ---
def memory_sweeper():
    if 'last_cleanup' not in st.session_state:
        st.session_state.last_cleanup = time.time()
        return
    if time.time() - st.session_state.last_cleanup > 300:
        st.session_state.word_buffer = None 
        gc.collect()
        st.session_state.last_cleanup = time.time()
    
    if 'last_hard_cleanup' not in st.session_state:
        st.session_state.last_hard_cleanup = time.time()
        
    if time.time() - st.session_state.last_hard_cleanup > 14400: # 4h
        st.cache_data.clear()
        gc.collect()
        st.session_state.last_hard_cleanup = time.time()

# --- LÓGICA DE FILA VISUAL ---
def get_ordered_visual_queue(queue, status_dict):
    if not queue: return []
    current_holder = next((c for c, s in status_dict.items() if 'Bastão' in (s or '')), None)
    if not current_holder or current_holder not in queue: return list(queue)
    try:
        idx = queue.index(current_holder)
        return queue[idx:] + queue[:idx]
    except ValueError: return list(queue)

# ============================================
# 5. MANIPULAÇÃO DE ESTADO E DATA
# ============================================

def reset_day_state():
    st.session_state.bastao_queue = []
    st.session_state.status_texto = {n: 'Indisponível' for n in CONSULTORES}
    st.session_state.daily_logs = []
    st.session_state.report_last_run_date = get_brazil_time()

def ensure_daily_reset():
    """
    CORREÇÃO 3: Lógica robusta para não zerar a fila acidentalmente.
    Só zera se a data salva no banco for VÁLIDA e MENOR que hoje.
    """
    now_br = get_brazil_time()
    last_run = st.session_state.get('report_last_run_date')
    
    # Tentativa segura de ler a data
    if isinstance(last_run, str):
        try: 
            last_run_dt = datetime.fromisoformat(last_run).date()
        except: 
            # Se a data estiver corrompida, assumimos que é HOJE para não zerar a fila.
            last_run_dt = now_br.date() 
    elif isinstance(last_run, datetime):
        last_run_dt = last_run.date()
    else:
        # Se for None ou outro tipo, assume hoje.
        last_run_dt = now_br.date()

    if now_br.date() > last_run_dt:
        if st.session_state.daily_logs: 
            send_daily_report_to_webhook()
            full_state = {
                'date': now_br.isoformat(),
                'logs': st.session_state.daily_logs,
                'queue_final': st.session_state.bastao_queue
            }
            send_state_dump_webhook(full_state)
        reset_day_state()
        st.session_state.report_last_run_date = now_br # Atualiza data
        save_state()

def auto_manage_time():
    ensure_daily_reset()

# ============================================
# 6. LOGICA DO SISTEMA (DOCUMENTOS E WEBHOOKS)
# ============================================
def verificar_duplicidade_certidao(tipo, processo=None, data=None):
    sb = get_supabase()
    if not sb: return False
    try:
        query = sb.table("certidoes_registro").select("*")
        if tipo in ['Físico', 'Eletrônico', 'Física', 'Eletrônica'] and processo:
            proc_limpo = str(processo).strip()
            if not proc_limpo: return False
            response = query.eq("processo", proc_limpo).execute()
            return len(response.data) > 0
        return False
    except Exception as e:
        return False

def salvar_certidao_db(dados):
    sb = get_supabase()
    if not sb: return False
    try:
        if isinstance(dados.get('data'), (date, datetime)):
            dados['data'] = dados['data'].isoformat()
        if 'hora_periodo' in dados:
             if dados['hora_periodo']:
                dados['motivo'] = f"{dados.get('motivo', '')} - Hora/Período: {dados['hora_periodo']}"
             del dados['hora_periodo']
        if 'n_processo' in dados: 
            dados['processo'] = dados.pop('n_processo')
        if 'n_chamado' in dados: 
            dados['incidente'] = dados.pop('n_chamado')
        if 'data_evento' in dados:
            dados['data'] = dados.pop('data_evento')

        sb.table("certidoes_registro").insert(dados).execute()
        return True
    except Exception as e:
        st.error(f"Erro ao salvar no Supabase: {e}")
        return False

def gerar_docx_certidao_internal(tipo, numero, data, consultor, motivo, chamado="", hora="", nome_parte=""):
    try:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0); section.right_margin = Cm(3.0)
        style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
        
        # Cabeçalho
        head_p = doc.add_paragraph(); head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runner = head_p.add_run("TRIBUNAL DE JUSTIÇA DO ESTADO DE MINAS GERAIS\n")
        runner.bold = True
        head_p.add_run("Rua Ouro Preto, N° 1564 - Bairro Santo Agostinho - CEP 30170-041 - Belo Horizonte - MG\nwww.tjmg.jus.br - Andar: 3º e 4º PV")
        doc.add_paragraph("\n")
        
        # Numeração e Assunto
        if tipo == 'Geral': p_num = doc.add_paragraph(f"Parecer GEJUD/DIRTEC/TJMG nº ____/2026. Assunto: Notifica erro no “JPe – 2ª Instância” ao peticionar.")
        else: p_num = doc.add_paragraph(f"Parecer Técnico GEJUD/DIRTEC/TJMG nº ____/2026. Assunto: Notifica erro no “JPe – 2ª Instância” ao peticionar.")
        p_num.runs[0].bold = True
        
        # Data
        data_extenso_str = ""
        try:
            dt_obj = datetime.strptime(data, "%d/%m/%Y")
            meses = {1:'janeiro', 2:'fevereiro', 3:'março', 4:'abril', 5:'maio', 6:'junho', 
                     7:'julho', 8:'agosto', 9:'setembro', 10:'outubro', 11:'novembro', 12:'dezembro'}
            data_extenso_str = f"Belo Horizonte, {dt_obj.day} de {meses[dt_obj.month]} de {dt_obj.year}"
        except:
            data_extenso_str = f"Belo Horizonte, {data}" 
            
        doc.add_paragraph(data_extenso_str)
              
        doc.add_paragraph(f"Exmo(a). Senhor(a) Relator(a),")
        
        if tipo == 'Geral':
            corpo = doc.add_paragraph(); corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            txt = (f"Para fins de cumprimento dos artigos 13 e 14 da Resolução nº 780/2014 do Tribunal de Justiça do Estado de Minas Gerais, "
                   f"informamos que em {data} houve indisponibilidade do portal JPe, superior a uma hora, {hora}, que impossibilitou o peticionamento eletrônico de recursos em processos que já tramitavam no sistema.")
            corpo.add_run(txt)
            doc.add_paragraph("\nColocamo-nos à disposição para outras que se fizerem necessárias.")
            
        elif tipo in ['Eletrônica', 'Eletrônico']:
            corpo = doc.add_paragraph(); corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo.add_run(f"Informamos que de {data}, houve indisponibilidade específica do sistema para o peticionamento do processo nº {numero}")
            if nome_parte: corpo.add_run(f", Parte/Advogado: {nome_parte}")
            corpo.add_run(".\n\n")
            corpo.add_run(f"O Chamado de número {chamado if chamado else '_____'}, foi aberto e encaminhado à DIRTEC (Diretoria Executiva de Tecnologia da Informação e Comunicação).\n\n")
            corpo.add_run("Esperamos ter prestado as informações solicitadas e colocamo-nos à disposição para outras que se fizerem necessárias.")

        elif tipo in ['Física', 'Físico']:
            corpo1 = doc.add_paragraph(); corpo1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo1.add_run(f"Informamos que no dia {data}, houve indisponibilidade específica do sistema para o peticionamento do processo nº {numero}")
            if nome_parte: corpo1.add_run(f", Parte/Advogado: {nome_parte}")
            corpo1.add_run(".")
            
            corpo2 = doc.add_paragraph(); corpo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo2.add_run(f"O Chamado de número {chamado if chamado else '_____'}, foi aberto e encaminhado à DIRTEC (Diretoria Executiva de Tecnologia da Informação e Comunicação).")
            
            corpo3 = doc.add_paragraph(); corpo3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            corpo3.add_run("Diante da indisponibilidade específica, não havendo um prazo para solução do problema, a Primeira Vice-Presidência recomenda o ingresso dos autos físicos, nos termos do § 2º, do artigo 14º, da Resolução nº 780/2014, do Tribunal de Justiça do Estado de Minas Gerais.")
            
            doc.add_paragraph("Colocamo-nos à disposição para outras informações que se fizerem necessárias.")

        doc.add_paragraph("\nRespeitosamente,")
        sign = doc.add_paragraph("\n___________________________________\nWaner Andrade Silva\n0-009020-9\nCoordenação de Análise e Integração de Sistemas Judiciais Informatizados - COJIN\nGerência de Sistemas Judiciais - GEJUD\nDiretoria Executiva de Tecnologia da Informação e Comunicação - DIRTEC")
        sign.runs[0].bold = True 
        
        buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
        return buffer
    except: return None

# Webhooks
def send_chat_notification_internal(consultor, status):
    """Notificação interna de giro corrigida para enviar os próximos."""
    if status != 'Bastão':
        return False
        
    # --- CORREÇÃO: Calcula os próximos em vez de mandar lista vazia ---
    lista_proximos = get_proximos_bastao(consultor, n=2)
    
    payload = {
        'evento': 'bastao_giro',
        'timestamp': get_brazil_time().isoformat(),
        'team_id': st.session_state.get('team_id'),
        'team_name': st.session_state.get('team_name'),
        'com_bastao_agora': consultor,
        'proximos': lista_proximos, # <--- AGORA VAI PREENCHIDO
    }
    return post_n8n(N8N_WEBHOOK_BASTAO_GIRO, payload)


def send_state_dump_webhook(state_data):
    if not WEBHOOK_STATE_DUMP: return False
    try:
        sanitized_data = clean_data_for_db(state_data)
        headers = {'Content-Type': 'application/json'}
        requests.post(WEBHOOK_STATE_DUMP, data=json.dumps(sanitized_data), headers=headers, timeout=5)
        return True
    except: return False

def send_horas_extras_to_chat(consultor, data, inicio, tempo, motivo):
    msg = f"⏰ **Registro de Horas Extras**\n\n👤 **Consultor:** {consultor}\n📅 **Data:** {data.strftime('%d/%m/%Y')}\n🕐 **Início:** {inicio.strftime('%H:%M')}\n⏱️ **Tempo Total:** {tempo}\n📝 **Motivo:** {motivo}"
    return notify_registro_ferramenta("H_EXTRAS", consultor, dados={"data": data.strftime("%d/%m/%Y"), "inicio": inicio.strftime("%H:%M"), "tempo": tempo, "motivo": motivo}, mensagem=msg)

def send_atendimento_to_chat(consultor, data, usuario, nome_setor, sistema, descricao, canal, desfecho, jira_opcional=""):
    jira_str = f"\n🔢 **Jira:** CESUPE-{jira_opcional}" if jira_opcional else ""
    msg = f"📋 **Novo Registro de Atendimento**\n\n👤 **Consultor:** {consultor}\n📅 **Data:** {data.strftime('%d/%m/%Y')}\n👥 **Usuário:** {usuario}\n🏢 **Nome/Setor:** {nome_setor}\n💻 **Sistema:** {sistema}\n📝 **Descrição:** {descricao}\n📞 **Canal:** {canal}\n✅ **Desfecho:** {desfecho}{jira_str}"
    return notify_registro_ferramenta("ATENDIMENTOS", consultor, dados={"data": data.strftime("%d/%m/%Y"), "usuario": usuario, "setor": nome_setor, "sistema": sistema, "canal": canal, "desfecho": desfecho, "jira": jira_opcional}, mensagem=msg)

def send_chamado_to_chat(consultor, texto):
    if not consultor or consultor == 'Selecione um nome' or not texto.strip(): return False
    data_envio = get_brazil_time().strftime('%d/%m/%Y %H:%M')
    msg = f"🆘 **Rascunho de Chamado/Jira**\n📅 **Data:** {data_envio}\n\n👤 **Autor:** {consultor}\n\n📝 **Texto:**\n{texto}"
    return notify_registro_ferramenta('CHAMADOS', consultor, dados={'texto': texto}, mensagem=msg)

def handle_erro_novidade_submission(consultor, titulo, objetivo, relato, resultado):
    data_envio = get_brazil_time().strftime("%d/%m/%Y %H:%M")
    msg = f"🐛 **Novo Relato de Erro/Novidade**\n📅 **Data:** {data_envio}\n\n👤 **Autor:** {consultor}\n📌 **Título:** {titulo}\n\n🎯 **Objetivo:**\n{objetivo}\n\n🧪 **Relato:**\n{relato}\n\n🏁 **Resultado:**\n{resultado}"
    return notify_registro_ferramenta("ERRO_NOVIDADE", consultor, dados={"titulo": titulo, "objetivo": objetivo, "relato": relato, "resultado": resultado}, mensagem=msg)

def send_sessao_to_chat_fn(consultor, texto_mensagem):
    return True

def handle_sugestao_submission(consultor, texto):
    data_envio = get_brazil_time().strftime("%d/%m/%Y %H:%M")
    ip_usuario = get_remote_ip()
    msg = f"💡 **Nova Sugestão**\n📅 **Data:** {data_envio}\n👤 **Autor:** {consultor}\n🌐 **IP:** {ip_usuario}\n\n📝 **Sugestão:**\n{texto}"
    return notify_registro_ferramenta("H_EXTRAS", consultor, dados={"data": data.strftime("%d/%m/%Y"), "inicio": inicio.strftime("%H:%M"), "tempo": tempo, "motivo": motivo}, mensagem=msg)

# ============================================
# 7. FUNÇÕES DE ESTADO E LÓGICA
# ============================================

def save_state():
    try:
        tid = st.session_state.get('team_id')
        if not tid: return # Evita salvar se nao tiver ID
        
        last_run = st.session_state.report_last_run_date
        visual_queue_calculated = get_ordered_visual_queue(st.session_state.bastao_queue, st.session_state.status_texto)
        
        state_to_save = {
            'status_texto': st.session_state.status_texto, 'bastao_queue': st.session_state.bastao_queue,
            'visual_queue': visual_queue_calculated, 'skip_flags': st.session_state.skip_flags, 
            'current_status_starts': st.session_state.current_status_starts,
            'bastao_counts': st.session_state.bastao_counts, 'priority_return_queue': st.session_state.priority_return_queue,
            'bastao_start_time': st.session_state.bastao_start_time, 
            'report_last_run_date': last_run, 
            'rotation_gif_start_time': st.session_state.get('rotation_gif_start_time'), 
            'auxilio_ativo': st.session_state.get('auxilio_ativo', False), 
            'daily_logs': st.session_state.daily_logs, 
            'simon_ranking': st.session_state.get('simon_ranking', []),
            'previous_states': st.session_state.get('previous_states', {}),
            'quick_indicators': st.session_state.get('quick_indicators', {})
        }
        
        save_state_to_db(tid, state_to_save)
        
        load_state_from_db.clear()
        try:
            st.session_state['_toast_msg'] = '✅ Registro salvo.'
            st.session_state['_skip_db_sync_until'] = time.time() + 2.0
        except Exception: pass
        
    except Exception as e: print(f"Erro save: {e}")

def format_time_duration(duration):
    if not isinstance(duration, timedelta): return '--:--:--'
    s = int(duration.total_seconds()); h, s = divmod(s, 3600); m, s = divmod(s, 60)
    return f'{h:02}:{m:02}:{s:02}'

def sync_state_from_db():
    try:
        # Se salvei há menos de 3s, confio no meu local, não puxo do banco
        if time.time() - st.session_state.get('_last_save_time', 0) < 3.0:
            return

        tid = st.session_state.get('team_id')
        if not tid: return

        db_data = load_state_from_db(tid)
        if not db_data: return

        keys = ['status_texto', 'bastao_queue', 'skip_flags', 'bastao_counts', 'priority_return_queue', 'daily_logs', 'simon_ranking', 'previous_states', 'quick_indicators']
        for k in keys:
            if k in db_data: 
                if k == 'daily_logs' and isinstance(db_data[k], list) and len(db_data[k]) > 150:
                    st.session_state[k] = db_data[k][-150:] 
                else:
                    st.session_state[k] = db_data[k]
                    
        if 'bastao_start_time' in db_data and db_data['bastao_start_time']:
            try:
                if isinstance(db_data['bastao_start_time'], str): st.session_state['bastao_start_time'] = datetime.fromisoformat(db_data['bastao_start_time'])
                else: st.session_state['bastao_start_time'] = db_data['bastao_start_time']
            except: pass
        if 'current_status_starts' in db_data:
            starts = db_data['current_status_starts']
            for nome, val in starts.items():
                if isinstance(val, str):
                    try: st.session_state.current_status_starts[nome] = datetime.fromisoformat(val)
                    except: pass
                else: st.session_state.current_status_starts[nome] = val
    except Exception as e: print(f"Erro sync: {e}")

def log_status_change(consultor, old_status, new_status, duration):
    if not isinstance(duration, timedelta): duration = timedelta(0)
    now_br = get_brazil_time()
    st.session_state.daily_logs.append({
        'timestamp': now_br, 'consultor': consultor, 
        'old_status': old_lbl if 'old_lbl' in locals() else old_status or 'Fila', 
        'new_status': new_lbl if 'new_lbl' in locals() else new_status or 'Fila', 
        'duration': duration, 'ip': st.session_state.get('device_id_val', 'unknown')
    })
    if len(st.session_state.daily_logs) > 150:
        st.session_state.daily_logs = st.session_state.daily_logs[-150:]
    st.session_state.current_status_starts[consultor] = now_br

# --- LOGICA DE STATUS (FUNÇÕES HELPERS) ---
def find_next_holder_index(current_index, queue, skips):
    if not queue: return -1
    n = len(queue); start_index = (current_index + 1) % n
    for i in range(n):
        idx = (start_index + i) % n
        if not skips.get(queue[idx], False): return idx
    if n > 1:
        proximo_imediato_idx = (current_index + 1) % n
        nome_escolhido = queue[proximo_imediato_idx]; st.session_state.skip_flags[nome_escolhido] = False 
        return proximo_imediato_idx
    return -1

def send_daily_report_to_webhook():
    logs = st.session_state.daily_logs
    if not logs: return False
    try: send_state_dump_webhook({'logs': st.session_state.daily_logs})
    except: pass

def check_and_assume_baton(forced_successor=None, immune_consultant=None):
    queue, skips = st.session_state.bastao_queue, st.session_state.skip_flags
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'Bastão' in s), None)
    is_valid = (current_holder and current_holder in queue)
    target = forced_successor if forced_successor else (current_holder if is_valid else None)
    
    if not target:
        curr_idx = queue.index(current_holder) if (current_holder and current_holder in queue) else -1
        idx = find_next_holder_index(curr_idx, queue, skips)
        target = queue[idx] if idx != -1 else None
        
    changed = False; now = get_brazil_time()
    
    for c in CONSULTORES:
        if c != immune_consultant: 
            if c != target and 'Bastão' in st.session_state.status_texto.get(c, ''):
                log_status_change(c, 'Bastão', 'Indisponível', now - st.session_state.current_status_starts.get(c, now))
                st.session_state.status_texto[c] = 'Indisponível'; changed = True
    
    if target:
        curr_s = st.session_state.status_texto.get(target, '')
        if 'Bastão' not in curr_s:
            old_s = curr_s; new_s = f"Bastão | {old_s}" if old_s and old_s != "Indisponível" else "Bastão"
            log_status_change(target, old_s, new_s, now - st.session_state.current_status_starts.get(target, now))
            st.session_state.status_texto[target] = new_s; st.session_state.bastao_start_time = now
            if current_holder != target: st.session_state.play_sound = True; send_chat_notification_internal(target, 'Bastão')
            st.session_state.skip_flags[target] = False
            changed = True
    elif not target and current_holder:
        if current_holder != immune_consultant:
            log_status_change(current_holder, 'Bastão', 'Indisponível', now - st.session_state.current_status_starts.get(current_holder, now))
            st.session_state.status_texto[current_holder] = 'Indisponível'; changed = True
            
    if changed: save_state()
    return changed

# --- AÇÕES PRINCIPAIS (QUE CHAMAM AS HELPERS) ---
def update_status(novo_status: str, marcar_indisponivel: bool = False, manter_fila_atual: bool = False):
    selected = st.session_state.get('consultor_selectbox')
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    
    ensure_daily_reset()
    now_br = get_brazil_time()
    current = st.session_state.status_texto.get(selected, '')
    forced_successor = None
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'Bastão' in (s or '')), None)
    
    if novo_status == 'Almoço':
        st.session_state.previous_states[selected] = {
            'status': current,
            'in_queue': selected in st.session_state.bastao_queue
        }
    
    if marcar_indisponivel:
        st.session_state[f'check_{selected}'] = False; st.session_state.skip_flags[selected] = True
        if selected in st.session_state.bastao_queue:
            if selected == current_holder:
                idx = st.session_state.bastao_queue.index(selected)
                nxt = find_next_holder_index(idx, st.session_state.bastao_queue, st.session_state.skip_flags)
                if nxt != -1: forced_successor = st.session_state.bastao_queue[nxt]
            st.session_state.bastao_queue.remove(selected)
            if selected not in st.session_state.priority_return_queue: st.session_state.priority_return_queue.append(selected)
    
    if not marcar_indisponivel:
        if novo_status == 'Indisponível':
            st.session_state[f'check_{selected}'] = False
            st.session_state.skip_flags[selected] = True
            if selected in st.session_state.bastao_queue:
                if selected == current_holder:
                    idx = st.session_state.bastao_queue.index(selected)
                    nxt = find_next_holder_index(idx, st.session_state.bastao_queue, st.session_state.skip_flags)
                    if nxt != -1: forced_successor = st.session_state.bastao_queue[nxt]
                st.session_state.bastao_queue.remove(selected)

        elif manter_fila_atual:
            if selected not in st.session_state.bastao_queue: st.session_state.bastao_queue.append(selected)
            st.session_state[f'check_{selected}'] = True
            st.session_state.skip_flags[selected] = False

        else:
            st.session_state[f'check_{selected}'] = False
            if selected in st.session_state.bastao_queue:
                if selected == current_holder:
                    idx = st.session_state.bastao_queue.index(selected)
                    nxt = find_next_holder_index(idx, st.session_state.bastao_queue, st.session_state.skip_flags)
                    if nxt != -1: forced_successor = st.session_state.bastao_queue[nxt]
                st.session_state.bastao_queue.remove(selected)
                if selected not in st.session_state.priority_return_queue: st.session_state.priority_return_queue.append(selected)
            st.session_state.skip_flags[selected] = True
            
    final_status = (novo_status or '').strip()
    if selected == current_holder and selected in st.session_state.bastao_queue:
        final_status = ('Bastão | ' + final_status).strip(' |') if final_status else 'Bastão'
    if not final_status and (selected not in st.session_state.bastao_queue): final_status = 'Indisponível'
    
    log_status_change(selected, current, final_status, now_br - st.session_state.current_status_starts.get(selected, now_br))
    
    st.session_state.status_texto[selected] = final_status
    check_and_assume_baton(forced_successor)
    save_state()


def get_bastao_holder_atual():
    """Retorna quem está com o Bastão (texto contém 'Bastão')."""
    return next((c for c, s in st.session_state.status_texto.items() if isinstance(s, str) and 'Bastão' in s), None)

def get_proximos_bastao(holder, n=3):
    """Retorna lista de próximos considerando skips."""
    queue = st.session_state.bastao_queue
    skips = st.session_state.skip_flags
    if not queue:
        return []
    if holder not in queue:
        # fallback: usa o primeiro da fila
        holder = queue[0]
    idx = queue.index(holder)
    proximos = []
    cursor = idx
    while len(proximos) < n:
        nxt = find_next_holder_index(cursor, queue, skips)
        if nxt == -1:
            break
        nxt_name = queue[nxt]
        # evita loop infinito
        if nxt_name == holder or nxt_name in proximos:
            break
        proximos.append(nxt_name)
        cursor = nxt
    return proximos

def notify_bastao_giro(reason='update', actor=None):
    """Envia para n8n quem está com o bastão e os próximos (silencioso)."""
    try:
        holder = get_bastao_holder_atual()
        if not holder and st.session_state.bastao_queue:
            holder = st.session_state.bastao_queue[0]
            
        lista_proximos = get_proximos_bastao(holder, n=2)
        txt_proximos = ", ".join(lista_proximos) if lista_proximos else "Ninguém"
        
        nome_equipe = st.session_state.get('team_name', 'Equipe')

        msg_final = (
            f"🔄 *Troca de Bastão - {nome_equipe}*\n\n"
            f"👤 *Agora:* {holder}\n"
            f"🔜 *Próximos:* {txt_proximos}"
        )

        payload = {
            'evento': 'bastao_giro',
            'motivo': reason,
            'timestamp': get_brazil_time().isoformat(),
            'team_id': st.session_state.get('team_id'),
            'team_name': nome_equipe,
            'actor': actor,
            'com_bastao_agora': holder,
            'proximos': lista_proximos,
            'tamanho_fila': len(st.session_state.bastao_queue),
            'message': msg_final  
        }
        
        post_n8n(N8N_WEBHOOK_BASTAO_GIRO, payload)
        return True
    except Exception:
        return False


def notify_registro_ferramenta(tipo: str, actor: str, dados: dict = None, mensagem: str = None) -> bool:
    """Envia evento de registro corrigindo o campo message."""
    if not mensagem:
        mensagem = f"Novo registro: {tipo} por {actor}"

    payload = {
        'evento': 'registro_ferramenta',
        'tipo': tipo,
        'timestamp': get_brazil_time().isoformat(),
        'team_id': st.session_state.get('team_id'),
        'team_name': st.session_state.get('team_name'),
        'actor': actor,
        'dados': dados or {},
        'message': mensagem, 
    }
    return post_n8n(N8N_WEBHOOK_REGISTROS, payload)


def toggle_queue(consultor):
    ensure_daily_reset(); st.session_state.gif_warning = False; now_br = get_brazil_time()
    if consultor in st.session_state.bastao_queue:
        current_holder = next((c for c, s in st.session_state.status_texto.items() if 'Bastão' in s), None)
        forced_successor = None
        if consultor == current_holder:
            idx = st.session_state.bastao_queue.index(consultor)
            nxt = find_next_holder_index(idx, st.session_state.bastao_queue, st.session_state.skip_flags)
            if nxt != -1: forced_successor = st.session_state.bastao_queue[nxt]
        st.session_state.bastao_queue.remove(consultor)
        st.session_state.status_texto[consultor] = 'Indisponível'
        check_and_assume_baton(forced_successor)
    else:
        st.session_state.bastao_queue.append(consultor)
        st.session_state.skip_flags[consultor] = False
        if consultor in st.session_state.priority_return_queue: st.session_state.priority_return_queue.remove(consultor)
        st.session_state.status_texto[consultor] = ''
        check_and_assume_baton()
        notify_bastao_giro(reason='enter_bastao', actor=consultor)
    save_state()

def rotate_bastao():
    ensure_daily_reset(); selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    
    queue = st.session_state.bastao_queue; skips = st.session_state.skip_flags
    current_holder = next((c for c, s in st.session_state.status_texto.items() if 'Bastão' in s), None)
    
    if selected != current_holder:
        st.error(f"⚠️ Apenas quem está com o bastão ({current_holder}) pode passá-lo!")
        return
        
    current_index = queue.index(current_holder) if current_holder in queue else -1
    if current_index == -1: check_and_assume_baton(); return
    next_idx = find_next_holder_index(current_index, queue, skips)
    if next_idx == -1 and len(queue) > 1: next_idx = (current_index + 1) % len(queue)
    if next_idx != -1:
        n_queue = len(queue); tmp_idx = (current_index + 1) % n_queue
        while tmp_idx != next_idx:
            skipped_name = queue[tmp_idx]
            if st.session_state.skip_flags.get(skipped_name, False): st.session_state.skip_flags[skipped_name] = False
            tmp_idx = (tmp_idx + 1) % n_queue
        next_holder = queue[next_idx]; st.session_state.skip_flags[next_holder] = False; now_br = get_brazil_time()
        old_h_status = st.session_state.status_texto[current_holder]
        new_h_status = old_h_status.replace('Bastão | ', '').replace('Bastão', '').strip()
        log_status_change(current_holder, old_h_status, new_h_status, now_br - (st.session_state.bastao_start_time or now_br))
        st.session_state.status_texto[current_holder] = new_h_status
        old_n_status = st.session_state.status_texto.get(next_holder, '')
        new_n_status = f"Bastão | {old_n_status}" if old_n_status else "Bastão"
        log_status_change(next_holder, old_n_status, new_n_status, timedelta(0))
        st.session_state.status_texto[next_holder] = new_n_status
        st.session_state.bastao_start_time = now_br
        st.session_state.bastao_counts[current_holder] = st.session_state.bastao_counts.get(current_holder, 0) + 1
        st.session_state.play_sound = True; send_chat_notification_internal(next_holder, 'Bastão')
        save_state()
    else: st.warning('Ninguém elegível.'); check_and_assume_baton()

def toggle_skip():
    selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    if selected not in st.session_state.bastao_queue: st.warning(f'{selected} não está na fila do bastão.'); return
    novo = not st.session_state.skip_flags.get(selected, False)
    st.session_state.skip_flags[selected] = novo
    save_state()

def toggle_presence_btn():
    selected = st.session_state.consultor_selectbox
    if not selected or selected == 'Selecione um nome': st.warning('Selecione um(a) consultor(a).'); return
    toggle_queue(selected)

# --- INDICADORES RÁPIDOS (📞 / ☕) ---
def _get_quick_indic(nome: str) -> dict:
    if 'quick_indicators' not in st.session_state or not isinstance(st.session_state.quick_indicators, dict):
        st.session_state.quick_indicators = {}
    return st.session_state.quick_indicators.get(nome, {'telefone': False, 'cafe': False})

def _set_quick_indic(nome: str, telefone=None, cafe=None):
    if 'quick_indicators' not in st.session_state or not isinstance(st.session_state.quick_indicators, dict):
        st.session_state.quick_indicators = {}
    cur = _get_quick_indic(nome).copy()
    if telefone is not None: cur['telefone'] = bool(telefone)
    if cafe is not None: cur['cafe'] = bool(cafe)
    st.session_state.quick_indicators[nome] = cur
    save_state()

def toggle_quick_telefone():
    nome = st.session_state.get('consultor_selectbox')
    if not nome or nome == 'Selecione um nome':
        st.warning('Selecione um(a) consultor(a).')
        return
    cur = _get_quick_indic(nome)
    novo = not cur.get('telefone', False)
    _set_quick_indic(nome, telefone=novo, cafe=False if novo else cur.get('cafe', False))

def toggle_quick_cafe():
    nome = st.session_state.get('consultor_selectbox')
    if not nome or nome == 'Selecione um nome':
        st.warning('Selecione um(a) consultor(a).')
        return
    cur = _get_quick_indic(nome)
    novo = not cur.get('cafe', False)
    _set_quick_indic(nome, cafe=novo, telefone=False if novo else cur.get('telefone', False))

def render_quick_toggle_btn(tipo: str):
    nome = st.session_state.get('consultor_selectbox')
    if not nome or nome == 'Selecione um nome':
        label = '📞' if tipo == 'telefone' else '☕'
        st.button(label, disabled=True, use_container_width=True)
        return
    
    indic = _get_quick_indic(nome)
    
    if tipo == 'telefone':
        ativo = bool(indic.get('telefone'))
        label = '📞✅' if ativo else '📞'
        if st.button(label, key=f'btn_tel_{nome}', use_container_width=True):
            toggle_quick_telefone(); st.rerun()
            
    elif tipo == 'cafe':
        ativo = bool(indic.get('cafe'))
        label = '☕✅' if ativo else '☕'
        if st.button(label, key=f'btn_cafe_{nome}', use_container_width=True):
            toggle_quick_cafe(); st.rerun()

def sync_logged_user():
    val = st.session_state.get('consultor_selectbox')
    if val and val != 'Selecione um nome':
        st.session_state['consultor_logado'] = val

def handle_sair():
    selected = st.session_state.get('consultor_selectbox')
    if not selected or selected == 'Selecione um nome':
        st.warning('Selecione um(a) consultor(a).')
        return
    ensure_daily_reset()
    current = st.session_state.status_texto.get(selected, '') or ''
    
    if current and current != 'Indisponível' and 'Bastão' not in current and current.strip() != '':
        if selected not in st.session_state.bastao_queue:
            st.session_state.bastao_queue.append(selected)
        st.session_state.status_texto[selected] = ''
        st.session_state.skip_flags[selected] = False
        if selected in st.session_state.priority_return_queue:
            st.session_state.priority_return_queue.remove(selected)
        check_and_assume_baton()
        save_state()
        return
    if selected in st.session_state.bastao_queue:
        toggle_queue(selected)
    else:
        st.session_state.status_texto[selected] = 'Indisponível'
        save_state()

def restore_from_lunch(nome: str):
    prev = st.session_state.get('previous_states', {}).get(nome)
    if not prev:
        st.session_state.status_texto[nome] = ''
        if nome not in st.session_state.bastao_queue:
            st.session_state.bastao_queue.append(nome)
        st.session_state.skip_flags[nome] = False
        check_and_assume_baton()
        save_state()
        return
    prev_status = prev.get('status', '') or ''
    prev_in_queue = bool(prev.get('in_queue', False))
    if isinstance(prev_status, str) and 'Bastão' in prev_status:
        prev_status = ''
        prev_in_queue = True
    st.session_state.status_texto[nome] = prev_status if prev_status != 'Indisponível' else ''
    if prev_in_queue:
        if nome not in st.session_state.bastao_queue:
            st.session_state.bastao_queue.append(nome)
        st.session_state.skip_flags[nome] = False
        if nome in st.session_state.priority_return_queue:
            st.session_state.priority_return_queue.remove(nome)
        check_and_assume_baton()
    else:
        if nome in st.session_state.bastao_queue:
            st.session_state.bastao_queue.remove(nome)
    try:
        del st.session_state.previous_states[nome]
    except Exception:
        pass
    save_state()

def handle_almoco_toggle():
    selected = st.session_state.get('consultor_selectbox')
    if not selected or selected == 'Selecione um nome':
        st.warning('Selecione um(a) consultor(a).')
        return
    current = st.session_state.status_texto.get(selected, '') or ''
    if current == 'Almoço':
        restore_from_lunch(selected)
        return
    update_status('Almoço', True)

def enter_from_indisponivel(c):
    if c not in st.session_state.bastao_queue: st.session_state.bastao_queue.append(c)
    st.session_state.status_texto[c] = ''; save_state()

def toggle_view(v):
    if st.session_state.active_view == v: st.session_state.active_view = None
    else: st.session_state.active_view = v

def init_session_state():
    dev = get_browser_id()
    if dev: 
        st.session_state['device_id_val'] = dev
    
    if 'db_loaded' not in st.session_state:
        tid = st.session_state.get('team_id', 2)
        db = load_state_from_db(tid)
        if db:
            if 'report_last_run_date' in db and isinstance(db['report_last_run_date'], str):
                try: 
                    db['report_last_run_date'] = datetime.fromisoformat(db['report_last_run_date'])
                except: 
                    db['report_last_run_date'] = datetime.min
            st.session_state.update(db)
        st.session_state['db_loaded'] = True
    
    defaults = {
        'bastao_start_time': None, 'report_last_run_date': datetime.min, 'rotation_gif_start_time': None,
        'play_sound': False, 'gif_warning': False, 'lunch_warning_info': None, 'last_reg_status': None,
        'chamado_guide_step': 0, 'auxilio_ativo': False, 'active_view': None,
        'consultor_selectbox': "Selecione um nome", 'status_texto': {n: 'Indisponível' for n in CONSULTORES},
        'bastao_queue': [], 'skip_flags': {}, 'current_status_starts': {n: get_brazil_time() for n in CONSULTORES},
        'bastao_counts': {n: 0 for n in CONSULTORES}, 'priority_return_queue': [], 'daily_logs': [], 'simon_ranking': [],
        'word_buffer': None, 'aviso_duplicidade': False, 'previous_states': {}, 'quick_indicators': {}, 'view_logmein_ui': False,
        'last_cleanup': time.time(), 'last_hard_cleanup': time.time()
    }
    for k, v in defaults.items():
        if k not in st.session_state: st.session_state[k] = v
    for n in CONSULTORES:
        st.session_state.skip_flags.setdefault(n, False)
        st.session_state[f'check_{n}'] = n in st.session_state.bastao_queue

def open_logmein_ui(): st.session_state.view_logmein_ui = True
def close_logmein_ui(): st.session_state.view_logmein_ui = False

# ============================================
# WATCHER INTELIGENTE (SEM CONFLITO DE SAVE)
# ============================================
# @st.fragment  # DESATIVADO (autorefresh full-run)
def watcher_de_atualizacoes():
    try:
        # marcador invisível para garantir execução periódica do fragment
        st.markdown("<div style='display:none'>refresh</div>", unsafe_allow_html=True)
        # Se salvei algo nos ultimos 5 segundos, NÃO atualize a tela agora.
        if time.time() - st.session_state.get('_last_save_time', 0) < 5.0:
            return

        tid = st.session_state.get('team_id', 2)
        sb = get_supabase()
        if not sb: return
        
        # 0) Verifica carimbo global (mudou em qualquer equipe => todos percebem)
        gver = load_global_state_version()
        seen = int(st.session_state.get('_global_state_version_seen') or 0)
        if gver and gver != seen:
            # Se o usuário estiver em formulário, evita perder preenchimento: adia refresh.
            if st.session_state.get('active_view'):
                st.session_state['_pending_global_refresh'] = True
                st.session_state['_global_state_version_seen'] = gver
                return
            st.session_state['_global_state_version_seen'] = gver
            try:
                load_state_from_db.clear()
                load_global_state_version.clear()
            except Exception:
                pass
            sync_state_from_db()
            st.rerun()

        # Se havia refresh pendente e não há mais formulário, executa agora
        if st.session_state.get('_pending_global_refresh') and not st.session_state.get('active_view'):
            st.session_state['_pending_global_refresh'] = False
            try:
                load_state_from_db.clear()
            except Exception:
                pass
            sync_state_from_db()
            st.rerun()

        # Leitura eficiente
        res = sb.table("app_state").select("data").eq("id", tid).execute()
        
        if res.data and len(res.data) > 0:
            remote_data = res.data[0].get("data", {})
            rem_status = remote_data.get('status_texto', {})
            rem_queue = remote_data.get('bastao_queue', [])
            loc_status = st.session_state.get('status_texto', {})
            loc_queue = st.session_state.get('bastao_queue', [])
            
            # Comparação direta de dicionários (mais segura que string)
            if rem_status != loc_status or rem_queue != loc_queue:
                load_state_from_db.clear()
                st.session_state.update(remote_data)
                st.rerun()
    except: pass

# ============================================
# PONTO DE ENTRADA
# ============================================
def render_dashboard(team_id: int, team_name: str, consultores_list: list, webhook_key: str, app_url: str, other_team_id: int, other_team_name: str, usuario_logado: str):
    # 1. Inicializa estado PRIMEIRO para garantir que tudo existe
    init_session_state()
    memory_sweeper()
    auto_manage_time()

    # 1.1) Garante que o team_id/other_team_id ficam disponíveis ANTES de qualquer sync
    if 'team_id' not in st.session_state or st.session_state.get('team_id') != team_id:
        st.session_state['team_id'] = team_id
        # Força recarregar logo na primeira vez que trocar de equipe
        try:
            load_state_from_db.clear()
        except Exception:
            pass

    st.session_state['team_name'] = team_name
    st.session_state['other_team_id'] = other_team_id
    st.session_state['other_team_name'] = other_team_name

    # 1.2) Variáveis globais
    global APP_URL_CLOUD, CONSULTORES
    APP_URL_CLOUD = app_url or APP_URL_CLOUD
    if consultores_list:
        CONSULTORES = list(consultores_list)

    # 2) Auto-refresh global (30s) — com "jitter" para não derrubar CPU quando vários usuários abrem ao mesmo tempo.
    #    (espalha os reruns no tempo e reduz picos)
    if st_autorefresh:
        st.session_state['_using_autorefresh_lib'] = True
        device_seed = str(st.session_state.get('device_id_val') or usuario_logado or team_id)
        jitter_ms = int(hashlib.sha256(device_seed.encode('utf-8')).hexdigest(), 16) % 4000  # 0..3999ms
        interval_ms = 30000 + jitter_ms  # ~30s (otimizado - economia de 33%)
        tick = st_autorefresh(interval=interval_ms, key=f"autorefresh_{team_id}")
        last_tick = st.session_state.get('_autorefresh_tick')
        pulse = (last_tick is None) or (tick != last_tick)
        st.session_state['_autorefresh_tick'] = tick
    else:
        st.session_state['_using_autorefresh_lib'] = False
        pulse = False

    # ========================================
    # TIMER EM TEMPO REAL (JavaScript)
    # ========================================
    try:
        # Verifica se auto-refresh está ativo
        if st.session_state.get("_using_autorefresh_lib"):
            st.markdown(
                '''
                <div id="timer-container" style="padding: 0.75rem 1.5rem; border-radius: 0.5rem; 
                     margin-bottom: 1rem; text-align: center; font-weight: bold; 
                     font-size: 1.1rem; transition: all 0.3s ease;">
                    <span id="timer-icon">🟢</span>
                    Próxima atualização automática em: 
                    <span id="timer-seconds" style="font-size: 1.3rem; font-weight: bold;">30</span>s
                </div>
                <script>
                (function() {
                    let seconds = 30;
                    const container = document.getElementById("timer-container");
                    const icon = document.getElementById("timer-icon");
                    const display = document.getElementById("timer-seconds");
                    
                    function updateTimer() {
                        if (seconds > 20) {
                            container.style.background = "linear-gradient(90deg, #10b981 0%, #10b98122 100%)";
                            container.style.color = "#1f2937";
                            icon.textContent = "🟢";
                        } else if (seconds > 10) {
                            container.style.background = "linear-gradient(90deg, #f59e0b 0%, #f59e0b22 100%)";
                            container.style.color = "#1f2937";
                            icon.textContent = "🟡";
                        } else {
                            container.style.background = "linear-gradient(90deg, #ef4444 0%, #ef444422 100%)";
                            container.style.color = "#1f2937";
                            icon.textContent = "🔴";
                        }
                        
                        display.textContent = seconds;
                        display.style.color = seconds > 20 ? "#10b981" : seconds > 10 ? "#f59e0b" : "#ef4444";
                        
                        if (seconds > 0) {
                            seconds--;
                        } else {
                            seconds = 30;
                        }
                    }
                    
                    updateTimer();
                    setInterval(updateTimer, 1000);
                })();
                </script>
                ''',
                unsafe_allow_html=True
            )
    except Exception as e:
        pass  # Timer é opcional, não quebra se falhar


    # 3) Sincronização: no pulso do autorefresh, puxa do banco.
    #    Se o usuário estiver com um "menu" aberto (active_view), não sobrescreve campos; marca pendente.
    if pulse:
        st.session_state['_last_autorefresh_ts'] = time.time()
        if st.session_state.get('active_view'):
            st.session_state['_pending_global_refresh'] = True
        else:
            st.session_state['_pending_global_refresh'] = False
            # Atualiza o estado da equipe atual (cache TTL já ajuda, mas aqui garantimos o "puxar" a cada pulso)
            sync_state_from_db()

    # Se tinha refresh pendente e o usuário fechou o menu, aplica já
    if st.session_state.get('_pending_global_refresh') and not st.session_state.get('active_view'):
        st.session_state['_pending_global_refresh'] = False
        sync_state_from_db()
    # global APP_URL_CLOUD  # já declarado acima
    try:
        msg_toast = st.session_state.get('_toast_msg')
        if msg_toast:
            try: st.toast(msg_toast)
            except: st.success(msg_toast)
            st.session_state['_toast_msg'] = None
    except: pass

    st.markdown("""<style>.sticky-topbar {position: sticky; top: 0; z-index: 999; background: rgba(255,255,255,0.98); border-bottom: 1px solid #eee; padding: 8px 12px; margin-bottom: 8px;} .sticky-topbar .muted { color: #666; font-size: 0.85rem; }</style>""", unsafe_allow_html=True)
    try:
        _user_top = st.session_state.get('consultor_logado') or st.session_state.get('consultor_selectbox') or '-'
        _team_top = st.session_state.get('team_name') or team_name or '-'
        _now_top = get_brazil_time().strftime('%d/%m/%Y %H:%M:%S')
        st.markdown(f"<div class='sticky-topbar'><div style='display:flex; justify-content:space-between; align-items:center; gap:12px;'><div><b>👤 Logado como:</b> {_user_top} &nbsp; <span class='muted'>|</span> &nbsp; <b>👥 Equipe:</b> {_team_top}</div><div class='muted'>🕒 {_now_top}</div></div></div>", unsafe_allow_html=True)
    except: pass

    st.session_state['team_id'] = team_id
    st.session_state['team_name'] = team_name
    st.session_state['other_team_id'] = other_team_id
    st.session_state['other_team_name'] = other_team_name
    st.session_state['webhook_key'] = webhook_key

    if usuario_logado:
        st.session_state['consultor_logado'] = usuario_logado
        if st.session_state.get('consultor_selectbox') in (None, '', 'Selecione um nome'):
            st.session_state['consultor_selectbox'] = usuario_logado

    st.markdown(
        """
<style>
/* ocupa melhor a largura (reduz “buraco” lateral) */
.block-container{padding-top:1rem !important; padding-left:1rem !important; padding-right:1rem !important; max-width: 1850px !important;}

/* Base dos botões (efeitos) */
div.stButton > button{
  width:100%;
  height:3rem;
  border-radius:14px !important;
  font-weight:800 !important;
  letter-spacing:.2px;
  transition: transform .12s ease, box-shadow .12s ease, filter .12s ease;
  box-shadow: 0 10px 24px rgba(0,0,0,0.08);
  user-select:none;
}
div.stButton > button:hover{
  transform: translateY(-1px);
  box-shadow: 0 0 0 2px rgba(255,255,255,0.40) inset, 0 14px 34px rgba(0,0,0,0.12);
  filter: brightness(1.04);
}
div.stButton > button:active{ transform: translateY(0px) scale(0.995); filter: brightness(0.98); }

/* Botão de menu (⋮) compacto e alinhado */
button[aria-label="⋮"]{
  width:100% !important;
  max-width:44px !important;
  min-width:34px !important;
  height:32px !important;
  padding:0 !important;
  border-radius:12px !important;
  font-weight:900 !important;
  margin-left:auto !important;
  margin-right:0 !important;
}
button[aria-label="⋮"]:hover{filter: brightness(0.98);}

/* Cores por botão (alvos por aria-label) */
button[aria-label="📋 Atividades"],
button[aria-label="🏗️ Projeto"],
button[aria-label="🎓 Treinamento"],
button[aria-label="📅 Reunião"],
button[aria-label="🍽️ Almoço"],
button[aria-label="🎙️ Sessão"],
button[aria-label="🚶 Saída"],
button[aria-label="🏃 Sair"],
button[aria-label="🤝 Atend. Presencial"],
button[aria-label="🎭 Entrar/Sair Fila"],
button[aria-label="🎯 Passar"],
button[aria-label="⏭️ Pular"]{
  color:#fff !important;
  border:0 !important;
  background-size: 220% 220%;
  background-position: 0% 50%;
}
button[aria-label="📋 Atividades"]{ background-image: linear-gradient(135deg, #22c55e, #16a34a); }
button[aria-label="🏗️ Projeto"]{ background-image: linear-gradient(135deg, #f59e0b, #f97316); }
button[aria-label="🎓 Treinamento"]{ background-image: linear-gradient(135deg, #8b5cf6, #a78bfa); }
button[aria-label="📅 Reunião"]{ background-image: linear-gradient(135deg, #3b82f6, #60a5fa); }
button[aria-label="🍽️ Almoço"]{ background-image: linear-gradient(135deg, #eab308, #facc15); color:#1f2937 !important; }
button[aria-label="🎙️ Sessão"]{ background-image: linear-gradient(135deg, #6366f1, #818cf8); }
button[aria-label="🚶 Saída"]{ background-image: linear-gradient(135deg, #64748b, #94a3b8); }
button[aria-label="🏃 Sair"]{ background-image: linear-gradient(135deg, #ef4444, #dc2626); }
button[aria-label="🤝 Atend. Presencial"]{ background-image: linear-gradient(135deg, #06b6d4, #22c55e); }
button[aria-label="🎭 Entrar/Sair Fila"]{ background-image: linear-gradient(135deg, #fb923c, #f97316); }
button[aria-label="🎯 Passar"]{ background-image: linear-gradient(135deg, #0ea5e9, #3b82f6); }
button[aria-label="⏭️ Pular"]{ background-image: linear-gradient(135deg, #94a3b8, #64748b); }

button[aria-label="📋 Atividades"]:hover,
button[aria-label="🏗️ Projeto"]:hover,
button[aria-label="🎓 Treinamento"]:hover,
button[aria-label="📅 Reunião"]:hover,
button[aria-label="🍽️ Almoço"]:hover,
button[aria-label="🎙️ Sessão"]:hover,
button[aria-label="🚶 Saída"]:hover,
button[aria-label="🏃 Sair"]:hover,
button[aria-label="🤝 Atend. Presencial"]:hover,
button[aria-label="🎭 Entrar/Sair Fila"]:hover,
button[aria-label="🎯 Passar"]:hover,
button[aria-label="⏭️ Pular"]:hover{
  background-position: 100% 50%;
  filter: brightness(1.03);
}


/* Grupos coloridos (envolva as linhas com <div class="btn-row ...">) */
.btn-row{ margin: 8px 0 10px 0; }

.btn-row.bastao div.stButton > button{
  color:#fff !important;
  border:0 !important;
  background-image: linear-gradient(135deg, rgba(255,140,0,.98), rgba(255,69,0,.96));
  background-size: 220% 220%;
  background-position: 0% 50%;
}
.btn-row.bastao div.stButton > button:hover{ background-position: 100% 50%; }

.btn-row.status div.stButton > button{
  color:#fff !important;
  border:0 !important;
  background-image: linear-gradient(135deg, rgba(30,136,229,.98), rgba(144,202,249,.92));
  background-size: 220% 220%;
  background-position: 0% 50%;
}
.btn-row.status div.stButton > button:hover{ background-position: 100% 50%; }

.btn-row.tools div.stButton > button{
  color: rgba(17,24,39,1) !important;
  border: 1px solid rgba(0,0,0,0.10) !important;
  background: linear-gradient(180deg, #ffffff, #f7f7f7);
}

/* Botões de perigo (Sair/Voltar) - continua vermelho, mas mais elegante */
button[aria-label="⬅️ SAIR / VOLTAR AO MENU"]{
  background: linear-gradient(135deg, #ef4444, #dc2626) !important;
  color: white !important;
  border: 0 !important;
  font-weight: 900 !important;
  height: 48px !important;
  border-radius: 14px !important;
  width: 100% !important;
}
button[aria-label="⬅️ SAIR / VOLTAR AO MENU"]:hover{filter: brightness(1.04);} 

</style>
        """,
        unsafe_allow_html=True,
    )
    st.components.v1.html("<script>window.scrollTo(0, 0);</script>", height=0)
    st.info("🎗️ Fevereiro Laranja: Apoie a conscientização sobre leucemia.")

    # --- HEADER FRAGMENT ---
    # @st.fragment  # DESATIVADO (autorefresh full-run)
    def render_header_info_left():
        sync_state_from_db()
        c_topo_esq, c_topo_dir = st.columns([2, 1], vertical_alignment="bottom")
        with c_topo_esq:
            img = get_img_as_base64_cached(PUG2026_FILENAME); src = f"data:image/png;base64,{img}" if img else GIF_BASTAO_HOLDER
            st.markdown(f"""<div style="display: flex; align-items: center; gap: 15px;"><h1 style="margin: 0; padding: 0; font-size: 2.2rem; color: #FF8C00; text-shadow: 1px 1px 2px #FF4500;">Controle Bastão Cesupe 2026 {BASTAO_EMOJI}</h1><img src="{src}" style="width: 150px; height: 150px; border-radius: 10px; border: 4px solid #FF8C00; object-fit: cover;"></div>""", unsafe_allow_html=True)
        with c_topo_dir:
            c_sub1, c_sub2 = st.columns([2, 1], vertical_alignment="bottom")
            with c_sub1: novo_responsavel = st.selectbox("Assumir Bastão (Rápido)", options=["Selecione"] + CONSULTORES, label_visibility="collapsed", key="quick_enter")
            with c_sub2:
                if st.button("🚀 Entrar", use_container_width=True, key="btn_entrar_header"):
                    if novo_responsavel != "Selecione": toggle_queue(novo_responsavel); st.rerun()
            st.caption(f"ID: ...{st.session_state.get('device_id_val', '???')[-4:]}")
            if st.button("🔄 Atualizar Agora", use_container_width=True): load_state_from_db.clear(); st.rerun()

        st.markdown("<hr style='border: 1px solid #FF8C00; margin-top: 5px; margin-bottom: 20px;'>", unsafe_allow_html=True)
        queue = st.session_state.bastao_queue
        skips = st.session_state.skip_flags
        responsavel = next((c for c, s in st.session_state.status_texto.items() if 'Bastão' in s), None)
        curr_idx = queue.index(responsavel) if responsavel in queue else -1
        prox_idx = find_next_holder_index(curr_idx, queue, skips)
        proximo = queue[prox_idx] if prox_idx != -1 else None

        st.header("Responsável pelo Bastão")
        if responsavel:
            st.markdown(f"""<div style="background: linear-gradient(135deg, #FFF3E0 0%, #FFFFFF 100%); border: 3px solid #FF8C00; padding: 25px; border-radius: 15px; display: flex; align-items: center; box-shadow: 0 4px 15px rgba(255, 140, 0, 0.3); margin-bottom: 20px;"><div style="flex-shrink: 0; margin-right: 25px;"><img src="{GIF_BASTAO_HOLDER}" style="width: 90px; height: 90px; border-radius: 50%; object-fit: cover; border: 2px solid #FF8C00;"></div><div><span style="font-size: 14px; color: #555; font-weight: bold; text-transform: uppercase; letter-spacing: 1.5px;">Atualmente com:</span><br><span style="font-size: 42px; font-weight: 800; color: #FF4500; line-height: 1.1;">{responsavel}</span></div></div>""", unsafe_allow_html=True)
            dur = get_brazil_time() - (st.session_state.bastao_start_time or get_brazil_time())
            st.caption(f"⏱️ Tempo com o bastão: **{format_time_duration(dur)}**")
        else: st.markdown('<h2>(Ninguém com o bastão)</h2>', unsafe_allow_html=True)
    
        st.markdown("###"); st.header("Próximos da Fila")
        if responsavel and responsavel in queue:
            c_idx = queue.index(responsavel)
            raw_ordered = queue[c_idx+1:] + queue[:c_idx]
        else: raw_ordered = list(queue)
        lista_pularam = [n for n in queue if skips.get(n, False) and n != responsavel]
        demais_na_fila = [n for n in raw_ordered if n != proximo and not skips.get(n, False)]
    
        if proximo:
            ic = _icons_telefone_cafe(st.session_state.get('quick_indicators', {}).get(proximo, {}))
            st.markdown(f"**Próximo Bastão:** {proximo}{ic}", unsafe_allow_html=True)
        else: st.markdown("**Próximo Bastão:** _Ninguém elegível_")

        if demais_na_fila:
            demais_fmt = [f"{n}{_icons_telefone_cafe(st.session_state.get('quick_indicators', {}).get(n, {}))}" for n in demais_na_fila]
            st.markdown("**Demais na fila:** " + ", ".join(demais_fmt), unsafe_allow_html=True)
        else: st.markdown("**Demais na fila:** _Vazio_")
        if lista_pularam: st.markdown(f"**Consultor(es) pulou(pularam) o bastão:** {', '.join(lista_pularam)}")

    # --- SIDEBAR FRAGMENT ---
    # @st.fragment  # DESATIVADO (autorefresh full-run)
    def render_right_sidebar():
        other_id = st.session_state.get('other_team_id')
        other_name = st.session_state.get('other_team_name', 'Outra Equipe')
        team_name = st.session_state.get('team_name', '')

        # Botão grande para voltar ao menu (tela de login)
        st.markdown(
            """
<style>
button[aria-label="⬅️ SAIR / VOLTAR AO MENU"]{
  background: #ef4444 !important;
  color: white !important;
  border: 1px solid #ef4444 !important;
  font-weight: 800 !important;
  height: 52px !important;
  border-radius: 14px !important;
  width: 100% !important;
}
button[aria-label="⬅️ SAIR / VOLTAR AO MENU"]:hover{
  background: #dc2626 !important;
  border-color: #dc2626 !important;
}
</style>
""",
            unsafe_allow_html=True
        )
        if st.button("⬅️ SAIR / VOLTAR AO MENU", use_container_width=True, key="btn_back_menu_top"):
            st.session_state['_force_back_to_names'] = True
            st.session_state['time_selecionado'] = None
            st.session_state['consultor_logado'] = None
            st.session_state['consultor_selectbox'] = 'Selecione um nome'
            st.rerun()

        with st.expander('🧭 Painel (outra equipe / LogMeIn / trocar consultor)', expanded=False):
            if st.button('🔙 Voltar à tela de nomes', use_container_width=True, key=f'btn_voltar_nomes_{uuid.uuid4().hex}'):
                st.session_state['_force_back_to_names'] = True
                st.session_state['time_selecionado'] = None
                st.session_state['consultor_logado'] = None
                st.session_state['consultor_selectbox'] = 'Selecione um nome'
                st.rerun()

            if other_id:
                try: other_state = load_state_from_db(other_id) or {}
                except: other_state = {}
                other_queue = other_state.get('bastao_queue', []) or []
                other_skips = other_state.get('skip_flags', {}) or {}
                other_status = other_state.get('status_texto', {}) or {}
                other_quick = other_state.get('quick_indicators', {}) or {}
                other_responsavel = None
                for c, s in (other_status or {}).items():
                    if isinstance(s, str) and 'Bastão' in s: other_responsavel = c; break
                proximo_outro = None
                if other_queue:
                    if other_responsavel in other_queue:
                        idx = other_queue.index(other_responsavel)
                        nxt = find_next_holder_index(idx, other_queue, other_skips)
                        proximo_outro = other_queue[nxt] if nxt != -1 else None
                    if not proximo_outro:
                        nxt = find_next_holder_index(-1, other_queue, other_skips)
                        proximo_outro = other_queue[nxt] if nxt != -1 else None
                def _fmt_other(nome):
                    rb = _badge_ramal_html(get_ramal_nome(nome))
                    ic = _icons_telefone_cafe(other_quick.get(nome, {}))
                    return f"{nome}{rb}{ic}"
                st.markdown(f"### 👥 Fila {other_name}")
                if other_responsavel: st.markdown(f"**Com o Bastão agora:** {_fmt_other(other_responsavel)}", unsafe_allow_html=True)
                else: st.markdown("**Com o Bastão agora:** _Ninguém_", unsafe_allow_html=True)
                if proximo_outro: st.markdown(f"**Próximo Bastão:** {_fmt_other(proximo_outro)}", unsafe_allow_html=True)
                else: st.markdown("**Próximo Bastão:** _Ninguém na fila_", unsafe_allow_html=True)
                if other_queue:
                    demais = [n for n in other_queue if n != other_responsavel and n != proximo_outro]
                    if demais:
                        st.markdown("**Demais na fila:**", unsafe_allow_html=True)
                        for n in demais: st.markdown(f"- {_fmt_other(n)}", unsafe_allow_html=True)
                    else: st.markdown("_Sem demais na fila._")
            else: st.markdown('_Sem outra equipe configurada._')

            st.divider(); st.markdown('### 🔑 LogMeIn'); st.caption('LogMeIn **não** é status: é apenas para visualizar quem está usando o programa.')
            c1, c2 = st.columns(2)
            with c1:
                if st.button('Abrir', use_container_width=True, key=f'btn_open_logmein_side_{uuid.uuid4().hex}'): open_logmein_ui()
            with c2:
                if st.button('Fechar', use_container_width=True, key=f'btn_close_logmein_side_{uuid.uuid4().hex}'): close_logmein_ui()

            if st.session_state.get('view_logmein_ui'):
                with st.container(border=True):
                    st.markdown('#### 💻 Acesso LogMeIn')
                    l_user, l_in_use = get_logmein_status()
                    st.image(GIF_LOGMEIN_TARGET, width=180)
                    if l_in_use:
                        st.error(f"🔴 EM USO POR: **{l_user}**")
                        meu_nome = st.session_state.get('consultor_selectbox')
                        if meu_nome == l_user or meu_nome in CONSULTORES:
                            if st.button('🔓 LIBERAR AGORA', type='primary', use_container_width=True, key=f'btn_logmein_liberar_side_{uuid.uuid4().hex}'):
                                set_logmein_status(None, False); close_logmein_ui(); st.rerun()
                        else: st.info('Aguarde a liberação.')
                    else:
                        st.success('✅ LIVRE PARA USO')
                        meu_nome = st.session_state.get('consultor_selectbox')
                        if meu_nome and meu_nome != 'Selecione um nome':
                            if st.button('🚀 ASSUMIR AGORA', use_container_width=True, key=f'btn_logmein_assumir_side_{uuid.uuid4().hex}'):
                                set_logmein_status(meu_nome, True); close_logmein_ui(); st.rerun()
                        else: st.warning('Selecione seu nome no topo para assumir.')
            st.divider()
            if isinstance(team_name, str) and 'eproc' in team_name.lower(): render_agenda_eproc_sidebar()

    def render_status_list():
        sync_state_from_db()
        queue = st.session_state.bastao_queue
        skips = st.session_state.skip_flags
        responsavel = next((c for c, s in st.session_state.status_texto.items() if 'Bastão' in s), None)

        st.header('Status dos(as) Consultores(as)')
        # --- Filtros (apenas visual) ---
        with st.expander("🔎 Filtros (não altera status)", expanded=False):
            filtro_exibir = st.selectbox("Exibir", ["Todos", "Só na fila", "Só fora da fila"], key="flt_exibir", index=0)
            opcoes_status = ["Na Fila", "Atend. Presencial", "Em Demanda", "Projetos", "Treinamento", "Reuniões", "Almoço", "Sessão", "Saída rápida", "Indisponível"]
            mostrar_status = st.multiselect("Mostrar", options=opcoes_status, default=opcoes_status, key="flt_mostrar_status")

        def _show_section(label: str) -> bool:
            if filtro_exibir == "Só na fila" and label != "Na Fila":
                return False
            if filtro_exibir == "Só fora da fila" and label == "Na Fila":
                return False
            return (label in mostrar_status)

        def _render_queue_menu(nome: str, key_base: str):
            """Menu (⋮) para entrar/sair da fila."""
            in_queue = (nome in st.session_state.bastao_queue)

            if hasattr(st, "popover"):
                try:
                    with st.popover("⋮"):
                        if in_queue:
                            if st.button("🚪 Sair da fila", key=f"{key_base}_out", use_container_width=True):
                                toggle_queue(nome)
                                st.session_state['_skip_db_sync_until'] = time.time() + 3
                                st.rerun()
                        else:
                            if st.button("➕ Entrar na fila", key=f"{key_base}_in", use_container_width=True):
                                toggle_queue(nome)
                                st.session_state['_skip_db_sync_until'] = time.time() + 3
                                st.rerun()
                    return
                except Exception:
                    pass

            # Fallback (sem popover): dropdown pequeno
            opts = ["⋮", "Entrar na fila", "Sair da fila"]
            choice_key = f"{key_base}_dd"
            choice = st.selectbox("", options=opts, key=choice_key, label_visibility="collapsed")
            if choice == "Entrar na fila" and not in_queue:
                st.session_state[choice_key] = "⋮"
                toggle_queue(nome)
                st.session_state['_skip_db_sync_until'] = time.time() + 3
                st.rerun()
            if choice == "Sair da fila" and in_queue:
                st.session_state[choice_key] = "⋮"
                toggle_queue(nome)
                st.session_state['_skip_db_sync_until'] = time.time() + 3
                st.rerun()

        ui_lists = {'fila': [], 'almoco': [], 'saida': [], 'ausente': [], 'atividade_especifica': [], 'sessao_especifica': [], 'projeto_especifico': [], 'reuniao_especifica': [], 'treinamento_especifico': [], 'indisponivel': [], 'presencial_especifico': []}
        for nome in CONSULTORES:
            if nome in st.session_state.bastao_queue: ui_lists['fila'].append(nome)
            status = st.session_state.status_texto.get(nome, 'Indisponível'); status = status if status is not None else 'Indisponível'
            if status in ('', None): pass
            elif status == 'Almoço': ui_lists['almoco'].append(nome)
            elif status == 'Saída rápida': ui_lists['saida'].append(nome)
            elif status == 'Indisponível' and nome not in st.session_state.bastao_queue: ui_lists['indisponivel'].append(nome)
            if isinstance(status, str):
                if 'Sessão:' in status or status.strip() == 'Sessão': ui_lists['sessao_especifica'].append((nome, status.replace('Sessão:', '').strip()))
                if 'Reunião:' in status or status.strip() == 'Reunião': ui_lists['reuniao_especifica'].append((nome, status.replace('Reunião:', '').strip()))
                if 'Projeto:' in status or status.strip() == 'Projeto': ui_lists['projeto_especifico'].append((nome, status.replace('Projeto:', '').strip()))
                if 'Treinamento:' in status or status.strip() == 'Treinamento': ui_lists['treinamento_especifico'].append((nome, status.replace('Treinamento:', '').strip()))
                if 'Atividade:' in status or status.strip() == 'Atendimento': ui_lists['atividade_especifica'].append((nome, status.replace('Atividade:', '').strip()))
                if 'Atendimento Presencial:' in status: ui_lists['presencial_especifico'].append((nome, status.replace('Atendimento Presencial:', '').strip()))

        if _show_section("Na Fila"):
            st.subheader(f'✅ Na Fila ({len(ui_lists["fila"])})')
            render_order = get_ordered_visual_queue(queue, st.session_state.status_texto)
            if not render_order and queue:
                render_order = list(queue)
            if not render_order:
                st.markdown('_Ninguém na fila._')
            else:
                for i, nome in enumerate(render_order, 1):
                    if nome not in ui_lists['fila']:
                        continue
                    col_nome, col_menu = st.columns([0.88, 0.12], vertical_alignment='center')
                    with col_menu:
                        _render_queue_menu(nome, key_base=f"fila_{nome}_frag")
                    skip_flag = skips.get(nome, False)
                    status_atual = st.session_state.status_texto.get(nome, '') or ''
                    extra = ''
                    indic_icons = _icons_telefone_cafe(st.session_state.get('quick_indicators', {}).get(nome, {}))
                    if 'Atividade' in status_atual:
                        extra += ' 📋'
                    if 'Projeto' in status_atual:
                        extra += ' 🏗️'
                    if nome == responsavel:
                        display = f'<span style="background-color: #FF8C00; color: #FFF; padding: 2px 6px; border-radius: 5px; font-weight: 800;">🎭 {i}º {nome}{indic_icons}</span>'
                    elif skip_flag:
                        display = f'<strong>{i}º {nome}{indic_icons}</strong>{extra} <span style="background-color: #FEF3C7; padding: 2px 8px; border-radius: 10px;">Pulando ⏭️</span>'
                    else:
                        display = f'<strong>{i}º {nome}{indic_icons}</strong>{extra} <span style="background-color: #FFEDD5; padding: 2px 8px; border-radius: 10px;">Aguardando</span>'
                    col_nome.markdown(display, unsafe_allow_html=True)
        else:
            st.markdown('_Seção Na Fila oculta pelos filtros._')

        st.markdown('---')

        def _render_section(label, titulo, icon, itens, cor, key_rm):
            if not _show_section(label):
                return
            colors = {'orange': '#FFF1E6', 'blue': '#E7F1FF', 'teal': '#E6FFFB', 'violet': '#F3E8FF', 'green': '#ECFDF3', 'red': '#FFE4E6', 'grey': '#F3F4F6', 'yellow': '#FEF9C3'}
            bg_hex = colors.get(cor, '#EEEEEE')
            st.subheader(f'{icon} {titulo} ({len(itens)})')
            if not itens:
                st.markdown(f'_Nenhum._')
            else:
                for item in itens:
                    nome = item[0] if isinstance(item, tuple) else item
                    desc = item[1] if isinstance(item, tuple) else titulo
                    col_n, col_menu = st.columns([0.88, 0.12], vertical_alignment='center')
                    with col_menu:
                        _render_queue_menu(nome, key_base=f"{key_rm}_{nome}")
                    indic_icons = _icons_telefone_cafe(st.session_state.get('quick_indicators', {}).get(nome, {}))
                    col_n.markdown(
                        f"<div style='font-size: 16px; margin: 2px 0;'><strong>{nome}{indic_icons}</strong><span style='background-color: {bg_hex}; color: #333; padding: 2px 8px; border-radius: 12px; font-size: 14px; margin-left: 8px;'>{desc}</span></div>",
                        unsafe_allow_html=True
                    )
            st.markdown('---')

        _render_section('Atend. Presencial', 'Atend. Presencial', '🤝', ui_lists['presencial_especifico'], 'yellow', 'presencial')
        _render_section('Em Demanda', 'Em Demanda', '📋', ui_lists['atividade_especifica'], 'orange', 'atividade')
        _render_section('Projetos', 'Projetos', '🏗️', ui_lists['projeto_especifico'], 'blue', 'projeto')
        _render_section('Treinamento', 'Treinamento', '🎓', ui_lists['treinamento_especifico'], 'teal', 'treinamento')
        _render_section('Reuniões', 'Reuniões', '📅', ui_lists['reuniao_especifica'], 'violet', 'reuniao')
        _render_section('Almoço', 'Almoço', '🍽️', ui_lists['almoco'], 'red', 'almoco')
        _render_section('Sessão', 'Sessão', '🎙️', ui_lists['sessao_especifica'], 'green', 'sessao')
        _render_section('Saída rápida', 'Saída rápida', '🚶', ui_lists['saida'], 'red', 'saida')
        _render_section('Indisponível', 'Indisponível', '❌', ui_lists['indisponivel'], 'grey', 'indisp')

    # =========================================================================
    # LAYOUT PRINCIPAL
    # =========================================================================
    col_principal, col_disponibilidade = st.columns([1.25, 1.05])

    with col_principal:
        render_header_info_left()

    with col_disponibilidade:
        render_right_sidebar()
        render_status_list()

    with col_principal:
        st.markdown("### 🎮 Painel de Ação")
        st.markdown('<div class="btn-row bastao">', unsafe_allow_html=True)
        st.markdown("**🎭 Ações do Bastão**")
        c_nome, c_act1, c_act2, c_act3 = st.columns([2, 1, 1, 1], vertical_alignment="bottom")
        with c_nome:
             st.caption(f"Logado como: **{st.session_state.get('consultor_logado', st.session_state.get('consultor_selectbox', ''))}**")
             sub_nome, sub_tel, sub_cafe = st.columns([3, 1.2, 1.2], vertical_alignment='bottom')
             with sub_nome: st.selectbox('Selecione:', ['Selecione um nome'] + CONSULTORES, key='consultor_selectbox', label_visibility='collapsed', on_change=sync_logged_user)
             with sub_tel: render_quick_toggle_btn('telefone')
             with sub_cafe: render_quick_toggle_btn('cafe')
        with c_act1:
            if st.button("🎭 Entrar/Sair Fila", use_container_width=True): toggle_presence_btn(); st.rerun()
        with c_act2:
            if st.button('🎯 Passar', use_container_width=True): rotate_bastao(); st.rerun()
        with c_act3:
            if st.button('⏭️ Pular', use_container_width=True): toggle_skip(); st.rerun()

        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown("**📌 Status**")
        st.markdown('<div class="btn-row status">', unsafe_allow_html=True)

        r2c1, r2c2, r2c3, r2c4, r2c5 = st.columns(5)
        if r2c1.button('📋 Atividades', use_container_width=True): toggle_view('menu_atividades'); st.rerun()
        if r2c2.button('🏗️ Projeto', use_container_width=True): toggle_view('menu_projetos'); st.rerun()
        if r2c3.button('🎓 Treinamento', use_container_width=True): toggle_view('menu_treinamento'); st.rerun()
        if r2c4.button('📅 Reunião', use_container_width=True): toggle_view('menu_reuniao'); st.rerun()
    
        almoco_label = '🍽️ Voltar' if (st.session_state.get('consultor_selectbox') and st.session_state.status_texto.get(st.session_state.get('consultor_selectbox'), '') == 'Almoço') else '🍽️ Almoço'
        if r2c5.button(almoco_label, use_container_width=True): handle_almoco_toggle(); st.rerun()
    
        r3c1, r3c2, r3c3, r3c4 = st.columns(4)
        if r3c1.button('🎙️ Sessão', use_container_width=True): toggle_view('menu_sessao'); st.rerun()
        if r3c2.button('🚶 Saída', use_container_width=True): update_status('Saída rápida', True); st.rerun()
        if r3c3.button('🏃 Sair', use_container_width=True): handle_sair(); st.rerun()
        if r3c4.button("🤝 Atend. Presencial", use_container_width=True): toggle_view('menu_presencial'); st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

        # --- MENUS DE AÇÃO ---
        if st.session_state.active_view == 'menu_atividades':
            with st.container(border=True):
                at_t = st.multiselect("Tipo:", OPCOES_ATIVIDADES_STATUS); at_e = st.text_input("Detalhe:")
                manter_bastao_ativ = st.checkbox("Continuar recebendo bastão? (Atividade)", value=False, key="keep_bastao_ativ")
                c1, c2, c3 = st.columns([1, 1, 1])
                with c1:
                    if st.button("Confirmar", type="primary", use_container_width=True): 
                        st.session_state.active_view = None
                        status_msg = f"Atividade: {', '.join(at_t)} - {at_e}"
                        if manter_bastao_ativ: update_status(status_msg, manter_fila_atual=True); st.rerun()
                        else: update_status(status_msg, marcar_indisponivel=True); st.rerun()
                with c2:
                    if st.button("Sair de atividades", use_container_width=True):
                        st.session_state.active_view = None; handle_sair(); st.rerun()
                with c3:
                    if st.button("Cancelar", use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == 'menu_presencial':
            with st.container(border=True):
                st.subheader('🤝 Registrar Atendimento Presencial'); local_presencial = st.text_input('Local:', key='pres_local'); objetivo_presencial = st.text_input('Objetivo:', key='pres_obj')
                c_ok, c_cancel = st.columns(2)
                with c_ok:
                    if st.button('✅ Confirmar', type='primary', use_container_width=True):
                        if not local_presencial.strip() or not objetivo_presencial.strip(): st.warning('Preencha Local e Objetivo.')
                        else: st.session_state.active_view = None; update_status(f"Atendimento Presencial: {local_presencial.strip()} - {objetivo_presencial.strip()}", True); st.rerun() 
                with c_cancel:
                    if st.button('❌ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == 'menu_projetos':
            with st.container(border=True):
                st.subheader('🏗️ Registrar Projeto')
                proj_nome = st.text_input('Nome do Projeto:', placeholder='Digite o nome do projeto...')
                manter_bastao = st.checkbox("Continuar recebendo bastão? (Modo Atividade)")
                c_ok, c_cancel = st.columns(2)
                with c_ok:
                    if st.button('✅ Confirmar', type='primary', use_container_width=True):
                        if not proj_nome.strip(): st.warning('Digite o nome do projeto.')
                        else: 
                            st.session_state.active_view = None
                            status_msg = f"Projeto: {proj_nome.strip()}"
                            if manter_bastao: update_status(status_msg, manter_fila_atual=True); st.rerun()
                            else: update_status(status_msg, marcar_indisponivel=True); st.rerun()
                with c_cancel:
                    if st.button('❌ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == 'menu_treinamento':
            with st.container(border=True):
                st.subheader('🎓 Registrar Treinamento'); tema = st.text_input('Tema/Conteúdo:'); obs = st.text_input('Observação (opcional):')
                c_ok, c_cancel = st.columns(2)
                with c_ok:
                    if st.button('✅ Confirmar', type='primary', use_container_width=True):
                        if not tema.strip(): st.warning('Informe o tema.')
                        else: st.session_state.active_view = None; update_status(f"Treinamento: {tema.strip()}" + (f" - {obs.strip()}" if obs.strip() else ""), True); st.rerun()
                with c_cancel:
                    if st.button('❌ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == 'menu_reuniao':
            with st.container(border=True):
                st.subheader('📅 Registrar Reunião'); assunto = st.text_input('Assunto:'); obs = st.text_input('Observação (opcional):')
                c_ok, c_cancel = st.columns(2)
                with c_ok:
                    if st.button('✅ Confirmar', type='primary', use_container_width=True):
                        if not assunto.strip(): st.warning('Informe o assunto.')
                        else: st.session_state.active_view = None; update_status(f"Reunião: {assunto.strip()}" + (f" - {obs.strip()}" if obs.strip() else ""), True); st.rerun()
                with c_cancel:
                    if st.button('❌ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == 'menu_sessao':
            with st.container(border=True):
                st.subheader('🎙️ Registrar Sessão')
                sessao_livre = st.text_input('Qual Sessão / Câmara?'); obs = st.text_input('Observação (opcional):')
                c_ok, c_cancel = st.columns(2)
                with c_ok:
                    if st.button('✅ Confirmar', type='primary', use_container_width=True):
                        consultor = st.session_state.get('consultor_selectbox')
                        if not consultor or consultor == 'Selecione um nome': st.error('Selecione um consultor.')
                        elif not sessao_livre.strip(): st.warning('Digite qual a sessão.')
                        else:
                            st.session_state.active_view = None; update_status(f"Sessão: {sessao_livre}" + (f" - {obs.strip()}" if obs.strip() else ""), True); st.rerun()
                with c_cancel:
                    if st.button('❌ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == "checklist":
            with st.container(border=True):
                st.header("Gerador de Checklist"); data_eproc = st.date_input("Data:", value=get_brazil_time().date()); camara_eproc = st.text_input("Câmara:")
                if st.button("Gerar HTML"): st.success("Checklist gerado!")
                if st.button("❌ Cancelar"): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == "chamados":
            with st.container(border=True):
                st.header('🆘 Chamados (Padrão / Jira)')
                st.text_area('Texto do chamado:', height=240, key='chamado_textarea')
                c1, c2 = st.columns(2)
                with c1:
                    if st.button('Enviar', type='primary', use_container_width=True):
                        if handle_chamado_submission(): st.success('Enviado!'); st.session_state.active_view = None; st.rerun()
                        else: st.error('Erro ao enviar.')
                with c2:
                        if st.button('❌ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == "atendimentos":
            with st.container(border=True):
                st.header('📝 Registro de Atendimentos')
                at_data = st.date_input('Data:', value=get_brazil_time().date())
                at_usuario = st.selectbox('Usuário:', REG_USUARIO_OPCOES); at_setor = st.text_input('Setor:'); at_sys = st.selectbox('Sistema:', REG_SISTEMA_OPCOES)
                at_desc = st.text_input('Descrição:'); at_canal = st.selectbox('Canal:', REG_CANAL_OPCOES); at_res = st.selectbox('Desfecho:', REG_DESFECHO_OPCOES); at_jira = st.text_input('Jira:')
                if st.button('Enviar', type='primary', use_container_width=True):
                    if send_atendimento_to_chat(st.session_state.consultor_selectbox, at_data, at_usuario, at_setor, at_sys, at_desc, at_canal, at_res, at_jira):
                        st.success('Enviado!'); st.session_state.active_view = None; st.rerun()
                    else: st.error('Erro.')
                if st.button('❌ Cancelar', use_container_width=True): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == "hextras":
            with st.container(border=True):
                st.header("⏰ Horas Extras"); d_ex = st.date_input("Data:"); h_in = st.time_input("Início:"); t_ex = st.text_input("Tempo Total:"); mot = st.text_input("Motivo:")
                if st.button("Registrar"): 
                    if send_horas_extras_to_chat(st.session_state.consultor_selectbox, d_ex, h_in, t_ex, mot): st.success("Registrado!"); st.session_state.active_view = None; st.rerun()
                if st.button("❌ Cancelar"): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == "erro_novidade":
            with st.container(border=True):
                st.header("🐛 Erro/Novidade"); tit = st.text_input("Título:"); obj = st.text_area("Objetivo:"); rel = st.text_area("Relato:"); res = st.text_area("Resultado:")
                if st.button("Enviar"): 
                    if handle_erro_novidade_submission(st.session_state.consultor_selectbox, tit, obj, rel, res): st.success("Enviado!"); st.session_state.active_view = None; st.rerun()
                if st.button("❌ Cancelar"): st.session_state.active_view = None; st.rerun()

        if st.session_state.active_view == "certidao":
            with st.container(border=True):
                st.header("🖨️ Registro de Certidão (2026)")
                c_data = st.date_input("Data do Evento:", value=get_brazil_time().date(), format="DD/MM/YYYY")
                tipo_cert = st.selectbox("Tipo:", ["Física", "Eletrônica", "Geral"])
                c_cons = st.session_state.consultor_selectbox
                c_hora = ""; c_motivo = st.text_area("Motivo/Detalhes:", height=100)
            
                if tipo_cert == "Geral": 
                    c_hora = st.text_input("Horário/Período (Ex: 13h às 15h):")
                    c_proc = ""; c_chamado = ""; c_nome_parte = ""; c_peticao = ""
                    if c_hora: c_motivo = f"{c_motivo} - Período: {c_hora}"
                else: 
                    c1, c2 = st.columns(2)
                    c_proc = c1.text_input("Processo (Com pontuação):")
                    c_chamado = c2.text_input("Incidente/Chamado:")
                    c3, c4 = st.columns(2)
                    c_nome_parte = c3.text_input("Nome da Parte/Advogado:")
                    c_peticao = c4.selectbox("Tipo de Petição:", ["Inicial", "Recursal", "Intermediária", "Outros"])
            
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("📄 Gerar Word", use_container_width=True): 
                        st.session_state.word_buffer = gerar_docx_certidao_internal(tipo_cert, c_proc, c_data.strftime("%d/%m/%Y"), c_cons, c_motivo, c_chamado, c_hora, c_nome_parte)
                    if st.session_state.word_buffer: st.download_button("⬇️ Baixar", st.session_state.word_buffer, file_name="certidao.docx")
                with c2:
                    if st.button("💾 Salvar e Notificar", type="primary", use_container_width=True):
                        if verificar_duplicidade_certidao(tipo_cert, c_proc, c_data): st.session_state.aviso_duplicidade = True
                        else:
                            payload = {"tipo": tipo_cert, "data": c_data.isoformat(), "consultor": c_cons, "incidente": c_chamado, "processo": c_proc, "motivo": c_motivo, "nome_parte": c_nome_parte, "peticao": c_peticao}
                            if salvar_certidao_db(payload):
                                msg_cert = f"🖨️ **Nova Certidão Registrada**\n👤 **Autor:** {c_cons}\n📅 **Data:** {c_data.strftime('%d/%m/%Y')}\n📄 **Tipo:** {tipo_cert}\n📂 **Proc:** {c_proc}"
                                notify_registro_ferramenta("CERTIDAO", st.session_state.consultor_selectbox, dados={"data": c_data.strftime("%d/%m/%Y"), "tipo": tipo_cert, "processo": c_proc}, mensagem=msg_cert)
                                st.success("Salvo!"); time.sleep(1); st.session_state.active_view = None; st.session_state.word_buffer = None; st.rerun()
                            else: st.error("Erro ao salvar no banco.")
                if st.button("❌ Cancelar"): st.session_state.active_view = None; st.rerun()
                if st.session_state.get('aviso_duplicidade'): st.error("⚠️ Este processo já possui registro de certidão!"); st.button("Ok, entendi", on_click=st.rerun)

        if st.session_state.active_view == "sugestao":
            with st.container(border=True):
                st.header("💡 Enviar Sugestão")
                sug_txt = st.text_area("Sua ideia ou melhoria:")
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("Enviar Sugestão", type="primary", use_container_width=True):
                        if handle_sugestao_submission(st.session_state.consultor_selectbox, sug_txt): st.success("Enviado com sucesso!"); st.session_state.active_view = None; st.rerun()
                        else: st.error("Erro ao enviar.")
                with c2:
                    if st.button("Cancelar", use_container_width=True): st.session_state.active_view = None; st.rerun()


        st.markdown("<hr style='border: 1px solid #FF8C00;'>", unsafe_allow_html=True)
        st.markdown("#### Ferramentas")
        st.markdown('<div class="btn-row tools">', unsafe_allow_html=True)
    
        c_t1, c_t2, c_t3, c_t4 = st.columns(4)
        c_t1.button("📑 Checklist", use_container_width=True, on_click=toggle_view, args=("checklist",))
        c_t2.button("🆘 Chamados", use_container_width=True, on_click=toggle_view, args=("chamados",))
        c_t3.button("📝 Atendimentos", use_container_width=True, on_click=toggle_view, args=("atendimentos",))
        c_t4.button("⏰ H. Extras", use_container_width=True, on_click=toggle_view, args=("hextras",))
    
        c_t5, c_t6, c_t7 = st.columns(3)
        c_t5.button("🐛 Erro/Novidade", use_container_width=True, on_click=toggle_view, args=("erro_novidade",))
        c_t6.button("🖨️ Certidão", use_container_width=True, on_click=toggle_view, args=("certidao",))
        c_t7.button("💡 Sugestão", use_container_width=True, on_click=toggle_view, args=("sugestao",))
        st.markdown('</div>', unsafe_allow_html=True)

        # CORREÇÃO 4: GRÁFICO MOVIDO PARA O FINAL (Fica abaixo de todos os formulários)
        with st.container(border=True):
            render_operational_summary()
